"""
novelforge.py - NovelForge v2.2.2

单文件完整实现。融合 InkOS 写作方法论、AI 痕迹检测、18 维度审计、写后验证。

变更日志：
  v2.2.2 - [Bug修复] 7项代码bug修复 + 2项结构性改进
         - [Fix] spot_fix 语法错误（括号位置）
         - [Fix] spot_fix Markdown列表规则误删对话行
         - [Fix] 角色名正则过宽，添加前后边界
         - [Fix] 情感弧线检测改为从字段提取而非全文扫描
         - [Fix] character_matrix 解析处理含"|"的字段
         - [Fix] 跨章检测改为读取台账而非完整正文
         - [Fix] 重复率计算分母改为 max
         - [New] 审计→修订循环：warning级审计问题触发自动修订
         - [New] 崩溃保护：记忆文件写入改为 .tmp + 原子替换
  v2.2.1 - 18维度审计系统 + 写后验证器 + 角色信息边界 + 支线停滞检测
  v2.2.0 - 融合 InkOS 核心写作方法论 + AI痕迹检测
  v2.1.0 - docx 排版输出
  v2.0.7 - 上下文精简、断点续写、checkpoint 管理

使用方式：
  python novelforge.py --generate-rules
  python novelforge.py 500
  python novelforge.py 500 201
  python novelforge.py 500 201 --rebuild

依赖：
  pip install python-docx
"""

import json
import re
import os
import math
import asyncio
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from collections import Counter

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ============================================================
# 配置
# ============================================================

CONFIG = {
    "config_dir": "config",
    "memory_dir": "memory",
    "output_dir": "output/chapters",
    "checkpoint_path": "checkpoint.json",
    "total_chapters": 500,
    "word_count_target": 3000,
    "max_rewrite_attempts": 3,
    "max_audit_revise_attempts": 2,
    "stale_subplot_threshold": 5,
    "repetition_lookback": 5,

    "docx": {
        "font_body": "宋体",
        "font_heading": "黑体",
        "size_body": 12,
        "size_heading": 15,
        "line_spacing": 1.5,
        "first_line_indent": 24,
        "margin_top": 2.54,
        "margin_bottom": 2.54,
        "margin_left": 3.18,
        "margin_right": 3.18,
    },
}


# ============================================================
# 工具函数
# ============================================================

def atomic_write(path: Path, content: str):
    """原子性写入：先写 .tmp，成功后 rename，防止崩溃导致文件损坏"""
    tmp_path = path.with_suffix('.tmp')
    tmp_path.write_text(content, encoding="utf-8")
    os.replace(str(tmp_path), str(path))


# ============================================================
# AI 痕迹检测（5 维度）
# ============================================================

HEDGE_WORDS = ["似乎", "可能", "或许", "大概", "某种程度上", "一定程度上", "在某种意义上"]
TRANSITION_WORDS = ["然而", "不过", "与此同时", "另一方面", "尽管如此", "话虽如此", "但值得注意的是"]
FATIGUE_WORDS = ["仿佛", "忽然", "竟然", "猛地", "猛然", "不禁", "宛如", "随即", "旋即", "霎时", "顿时"]


def analyze_ai_tells(content: str) -> List[Dict]:
    issues = []
    total_chars = len(content)

    # 维度 1：段落等长
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', content) if p.strip()]
    if len(paragraphs) >= 3:
        lengths = [len(p) for p in paragraphs]
        mean = sum(lengths) / len(lengths)
        if mean > 0:
            cv = math.sqrt(sum((l - mean) ** 2 for l in lengths) / len(lengths)) / mean
            if cv < 0.15:
                issues.append({"severity": "warning", "category": "段落等长",
                    "description": f"变异系数{cv:.3f}（阈值<0.15）",
                    "suggestion": "增加段落长度差异"})

    # 维度 2：套话密度
    if total_chars > 0:
        density = sum(len(re.findall(re.escape(w), content)) for w in HEDGE_WORDS) / (total_chars / 1000)
        if density > 3:
            issues.append({"severity": "warning", "category": "套话密度",
                "description": f"{density:.1f}次/千字（阈值>3）",
                "suggestion": "用确定性叙述替代模糊表达"})

    # 维度 3：公式化转折
    trans = {w: content.count(w) for w in TRANSITION_WORDS if content.count(w) >= 3}
    if trans:
        issues.append({"severity": "warning", "category": "公式化转折",
            "description": "、".join(f'"{w}"×{c}' for w, c in trans.items()),
            "suggestion": "换用动作切入、时间跳跃、视角切换"})

    # 维度 4：列表式结构
    sentences = [s.strip() for s in re.split(r'[。！？\n]', content) if len(s.strip()) > 2]
    if len(sentences) >= 3:
        mc, c = 1, 1
        for i in range(1, len(sentences)):
            if sentences[i - 1][:2] == sentences[i][:2]:
                c += 1; mc = max(mc, c)
            else:
                c = 1
        if mc >= 3:
            issues.append({"severity": "info", "category": "列表式结构",
                "description": f"{mc}句连续相同开头", "suggestion": "变换句式开头"})

    # 维度 5：高疲劳词
    if total_chars > 0:
        threshold = max(1, total_chars // 3000)
        hits = {w: content.count(w) for w in FATIGUE_WORDS if content.count(w) > threshold}
        if hits:
            issues.append({"severity": "warning", "category": "高疲劳词",
                "description": "、".join(f'"{w}"×{c}' for w, c in hits.items()),
                "suggestion": "改用具体动作或感官描写"})

    return issues


def format_ai_tell_report(issues: List[Dict]) -> str:
    if not issues:
        return "AI痕迹检测：通过"
    lines = ["AI痕迹检测：发现问题"]
    for i in issues:
        icon = "⚠" if i["severity"] == "warning" else "ℹ"
        lines.append(f"  {icon} [{i['category']}] {i['description']}")
        lines.append(f"    → {i['suggestion']}")
    return "\n".join(lines)


# ============================================================
# PostWriteValidator - 写后验证器
# ============================================================

class PostWriteValidator:

    FORBIDDEN_TERMS = [
        "核心动机", "信息边界", "信息落差", "核心风险", "利益最大化",
        "当前处境", "认知共鸣", "锚定效应", "沉没成本", "情绪缺口",
    ]

    def spot_fix(self, content: str) -> Tuple[str, List[str]]:
        """硬规则自动 spot-fix，返回 (修复后内容, 修复记录)"""
        fixes = []

        # 1. 破折号替换
        if "——" in content:
            content = content.replace("——", "，")
            fixes.append("破折号已替换为逗号")

        # 2. 分析报告术语标记
        for term in self.FORBIDDEN_TERMS:
            if term in content:
                content = content.replace(term, "...")
                fixes.append(f"术语'{term}'已标记")

        # 3. "不是...而是..."句式标记
        nb_matches = re.findall(r'不是[^，。！？\n]{1,20}[，,]而?是[^，。！？\n]{1,20}', content)
        if nb_matches:
            for m in nb_matches:
                content = content.replace(m, f"[待改写:{m}]")
            fixes.append(f"'不是而是'{len(nb_matches)}处已标记")

        # 4. 章节号指称
        ch_refs = set(re.findall(r'第\d+章|chapter\s*\d+', content, re.IGNORECASE))
        if ch_refs:
            for ref in ch_refs:
                content = content.replace(ref, "...")
            fixes.append(f"章节号指称{len(ch_refs)}处已清除")

        return content, fixes

    def check_repetition_with_summaries(
        self, chapter_num: int, content: str, summaries_text: str
    ) -> List[Dict]:
        """
        跨章重复检测：基于台账摘要而非完整正文。
        检测：关键词重叠、章内重复句。
        """
        issues = []
        lookback = CONFIG["repetition_lookback"]

        # 从台账提取最近几章的摘要
        recent_summaries = []
        for ch in range(max(1, chapter_num - lookback), chapter_num):
            pattern = rf'第{ch}章[^\n]*\n(.*?)(?=\n第\d+章|\n## |\Z)'
            m = re.search(pattern, summaries_text, re.DOTALL)
            if m:
                recent_summaries.append((ch, m.group(1).strip()))

        # 本章关键词
        current_words = set(re.findall(r'[\u4e00-\u9fff]{2,4}', content))
        # 过滤停用词
        stop = {"一个", "这个", "那个", "他们", "她们", "我们", "什么", "怎么", "可以", "没有", "但是", "然后", "因为", "所以", "如果"}
        current_words -= stop

        for ch, summary in recent_summaries:
            summary_words = set(re.findall(r'[\u4e00-\u9fff]{2,4}', summary)) - stop
            if current_words and summary_words:
                overlap = len(current_words & summary_words) / max(len(current_words), len(summary_words))
                if overlap > 0.5:
                    issues.append({
                        "severity": "warning", "category": "跨章词汇重复",
                        "description": f"与第{ch}章关键词重叠率{overlap:.1%}",
                        "suggestion": "检查是否有重复描写或情节",
                    })

        # 章内重复句子
        sentences = [s.strip() for s in re.split(r'[。！？]', content) if len(s.strip()) > 10]
        seen: Dict[str, int] = {}
        for s in sentences:
            key = s[:15]
            seen[key] = seen.get(key, 0) + 1
        dups = [(k, c) for k, c in seen.items() if c >= 2]
        if dups:
            issues.append({
                "severity": "warning", "category": "章内句子重复",
                "description": f"{len(dups)}组相似句子",
                "suggestion": "删除或改写重复句子",
            })

        return issues


# ============================================================
# AuditSystem - 18 维度审计
# ============================================================

class AuditSystem:

    # 可通过修订修复的审计类别
    FIXABLE_CATEGORIES = {"资源连续性", "信息越界", "违禁模式", "违禁句式", "违禁符号", "违禁术语",
                          "跨章词汇重复", "章内句子重复", "高疲劳词", "段落等长", "套话密度", "公式化转折"}

    def __init__(self, memory_dir: str, output_dir: str):
        self.memory_dir = Path(memory_dir)
        self.output_dir = Path(output_dir)
        self.validator = PostWriteValidator()

    # ---- 统一入口 ----

    def run_audit(self, chapter_num: int, content: str, memory: Dict[str, str]) -> List[Dict]:
        all_issues = []
        all_issues.extend(self._layer1(chapter_num, content, memory))
        if chapter_num % 5 == 0:
            all_issues.extend(self._layer2(chapter_num, content, memory))
        if chapter_num % 20 == 0:
            all_issues.extend(self._layer3(chapter_num, memory))
        return all_issues

    def get_fixable_warnings(self, issues: List[Dict]) -> List[Dict]:
        """提取可通过修订修复的 warning 级问题"""
        return [i for i in issues
                if i["severity"] == "warning" and i["category"] in self.FIXABLE_CATEGORIES]

    # ---- Layer 1：每章，纯代码（7 维度）----

    def _layer1(self, ch: int, content: str, memory: Dict[str, str]) -> List[Dict]:
        issues = []
        issues.extend(self._check_character_names(content, memory.get("character_tracking", "")))
        issues.extend(self._check_timeline(content, ch))
        issues.extend(self._check_resources(content, memory.get("resource_ledger", "")))
        issues.extend(self._check_info_boundary(content, memory.get("character_matrix", "")))
        issues.extend(self._check_forbidden_patterns(content))
        issues.extend(self.validator.check_repetition_with_summaries(
            ch, content, memory.get("chapter_summaries", "")))
        issues.extend(analyze_ai_tells(content))
        return issues

    # ---- Layer 2：每 5 章（4 维度）----

    def _layer2(self, ch: int, content: str, memory: Dict[str, str]) -> List[Dict]:
        issues = []
        issues.extend(self._check_foreshadowing(memory.get("foreshadowing", ""), ch))
        issues.extend(self._check_subplot_stagnation(memory.get("plot_threads", ""), ch))
        issues.extend(self._check_emotional_consistency(memory.get("emotional_arcs", "")))
        issues.extend(self._check_pacing(memory.get("chapter_summaries", ""), ch))
        return issues

    # ---- Layer 3：每 20 章（4 维度）----

    def _layer3(self, ch: int, memory: Dict[str, str]) -> List[Dict]:
        issues = []
        issues.extend(self._check_vocabulary_diversity(memory.get("chapter_summaries", ""), ch))
        issues.extend(self._check_sentence_patterns(memory.get("chapter_summaries", "")))
        issues.extend(self._check_character_balance(memory.get("chapter_summaries", ""), ch))
        issues.extend(self._check_numerical_consistency(memory.get("resource_ledger", "")))
        return issues

    # ---- 维度实现 ----

    def _check_character_names(self, content: str, tracking: str) -> List[Dict]:
        if not tracking:
            return []
        known = set()
        for line in tracking.split("\n"):
            m = re.match(r'-\s*(\S+)', line.strip())
            if m:
                known.add(m.group(1))
        for m in re.finditer(r'(\S{2,4}?)[\s|｜]', tracking):
            if len(m.group(1)) >= 2:
                known.add(m.group(1))
        if not known:
            return []

        # [Bug3修复] 添加前后边界，避免误匹配
        dialog_names = re.findall(r'["「](\S{2,4}?)[」"]\s*[,，。]?\s*(?:说|道|问|答|喊|笑|叹)', content)
        # 动作名：前面必须是标点或句首，后面紧跟动词
        action_names = re.findall(r'(?:^|[。！？\n，])\s*(\S{2,4}?)(?:说|道|问|答|走|坐|站|怒|惊)', content)

        all_names = set(dialog_names + action_names)
        unknown = all_names - known - {"他", "她", "我", "你", "我们", "他们", "对方", "那人", "这人", "一个", "突然", "忍不住", "忽然", "急忙", "赶紧"}
        if unknown:
            return [{"severity": "info", "category": "角色名一致性",
                "description": f"未在台账中的角色名：{'、'.join(list(unknown)[:5])}",
                "suggestion": "确认是否为新角色"}]
        return []

    def _check_timeline(self, content: str, ch: int) -> List[Dict]:
        issues = []
        for num, unit in re.findall(r'(\d+)\s*(天|日|月|年|小时|分钟)\s*(?:后|前|以后|以前)', content):
            n = int(num)
            if unit == "年" and n > 10:
                issues.append({"severity": "warning", "category": "时间线连续性",
                    "description": f"第{ch}章出现{n}{unit}的时间跳跃",
                    "suggestion": "过大的时间跳跃需要铺垫"})
            elif unit == "月" and n > 6:
                issues.append({"severity": "info", "category": "时间线连续性",
                    "description": f"第{ch}章出现{n}{unit}的时间跳跃",
                    "suggestion": "确认是否需要过渡"})
        return issues

    def _check_resources(self, content: str, ledger: str) -> List[Dict]:
        if not ledger:
            return []
        issues = []
        for m in re.finditer(r'(\S{2,8}?)\s*[:：]\s*(\d+)', ledger):
            name, expected = m.group(1), int(m.group(2))
            if name in content:
                actual_match = re.search(re.escape(name) + r'\s*[:：有剩]\s*(\d+)', content)
                if actual_match:
                    actual = int(actual_match.group(1))
                    if actual != expected:
                        issues.append({"severity": "warning", "category": "资源连续性",
                            "description": f"'{name}'账本{expected}，正文{actual}",
                            "suggestion": "核实资源数量"})
        return issues

    def _check_info_boundary(self, content: str, matrix: str) -> List[Dict]:
        if not matrix:
            return []
        issues = []
        for line in matrix.split("\n"):
            if any(kw in line for kw in ("不知道", "未知", "未被告知")):
                # [Bug5修复] 用限制性拆分，只取前3个字段
                parts = line.split("|", 2)
                if len(parts) >= 3:
                    unknown_info = parts[2].strip()
                    if len(unknown_info) >= 4 and unknown_info in content:
                        issues.append({"severity": "warning", "category": "信息越界",
                            "description": f"正文出现标记为'未知'的信息：'{unknown_info[:20]}'",
                            "suggestion": "检查信息越界"})
        return issues

    def _check_forbidden_patterns(self, content: str) -> List[Dict]:
        issues = []
        nb = re.findall(r'不是[^，。！？\n]{1,20}[，,]而?是[^，。！？\n]{1,20}', content)
        if nb:
            issues.append({"severity": "warning", "category": "违禁句式",
                "description": f"'不是而是'{len(nb)}处", "suggestion": "改用直述句"})
        dash_count = content.count("——")
        if dash_count > 0:
            issues.append({"severity": "warning", "category": "违禁符号",
                "description": f"破折号{dash_count}处", "suggestion": "用逗号或句号"})
        report_terms = ["核心动机", "信息边界", "信息落差", "核心风险", "利益最大化", "当前处境"]
        found = [t for t in report_terms if t in content]
        if found:
            issues.append({"severity": "warning", "category": "违禁术语",
                "description": f"分析术语：{'、'.join(found)}", "suggestion": "替换为口语化表达"})
        return issues

    def _check_foreshadowing(self, foreshadowing: str, ch: int) -> List[Dict]:
        if not foreshadowing:
            return []
        issues = []
        for line in foreshadowing.split("\n"):
            line = line.strip()
            if not line or "已回收" in line:
                continue
            ch_matches = re.findall(r'第(\d+)章', line)
            if ch_matches:
                gap = ch - max(int(c) for c in ch_matches)
                if gap >= 10:
                    issues.append({"severity": "warning", "category": "伏笔回收",
                        "description": f"'{line[:30]}'已{gap}章未推进",
                        "suggestion": "回收或推进该伏笔"})
        return issues

    def _check_subplot_stagnation(self, plot_threads: str, ch: int) -> List[Dict]:
        if not plot_threads:
            return []
        threshold = CONFIG["stale_subplot_threshold"]
        issues = []
        for line in plot_threads.split("\n"):
            line = line.strip()
            if not line or "已解决" in line or "已完成" in line:
                continue
            ch_matches = re.findall(r'第(\d+)章', line)
            if ch_matches:
                gap = ch - max(int(c) for c in ch_matches)
                if gap >= threshold:
                    issues.append({"severity": "warning", "category": "支线停滞",
                        "description": f"'{line[:30]}'已{gap}章未推进",
                        "suggestion": "推进或给出搁置理由"})
        return issues

    def _check_emotional_consistency(self, emotional_arcs: str) -> List[Dict]:
        if not emotional_arcs:
            return []
        positive = {"开心", "兴奋", "自信", "满足", "平静", "温暖", "希望", "释然"}
        negative = {"愤怒", "悲伤", "恐惧", "焦虑", "绝望", "痛苦", "压抑", "不安"}
        # [Bug4修复] 从情绪状态字段提取，而非全文扫描
        character_emotions: Dict[str, List[str]] = {}
        for line in emotional_arcs.split("\n"):
            line = line.strip()
            if not line or not line.startswith("-"):
                continue
            parts = re.split(r'[\|｜]', line.lstrip("- "))
            if len(parts) >= 2:
                name = parts[0].strip()
                emotion_field = parts[1].strip() if len(parts) >= 2 else ""
                # 从情绪字段（通常是第二个字段）提取
                emotion_words = [w for w in positive | negative if w in emotion_field]
                if emotion_words:
                    character_emotions.setdefault(name, []).append(emotion_words[0])

        issues = []
        for name, emotions in character_emotions.items():
            if len(emotions) >= 3:
                recent = emotions[-3:]
                flips = sum(1 for i in range(1, len(recent))
                           if (recent[i] in positive) != (recent[i - 1] in positive))
                if flips >= 2:
                    issues.append({"severity": "info", "category": "情感弧线",
                        "description": f"'{name}'情绪反复：{'→'.join(recent)}",
                        "suggestion": "情绪变化需要事件驱动"})
        return issues

    def _check_pacing(self, summaries: str, ch: int) -> List[Dict]:
        if not summaries:
            return []
        high = {"战斗", "冲突", "高潮", "追杀", "对决", "爆发", "危机", "对峙", "翻脸"}
        low = {"过渡", "日常", "修炼", "铺垫", "准备", "休息", "闲聊", "整理"}
        recent_types = []
        for line in summaries.split("\n"):
            m = re.search(r'第(\d+)章', line)
            if m:
                c = int(m.group(1))
                if ch - 5 <= c <= ch:
                    if any(w in line for w in high):
                        recent_types.append("high")
                    elif any(w in line for w in low):
                        recent_types.append("low")
        issues = []
        if len(recent_types) >= 4:
            tail = recent_types[-4:]
            if all(t == "high" for t in tail):
                issues.append({"severity": "info", "category": "节奏张弛",
                    "description": f"最近{len(tail)}章全是高张力",
                    "suggestion": "插入过渡章节"})
            elif all(t == "low" for t in tail):
                issues.append({"severity": "warning", "category": "节奏张弛",
                    "description": f"最近{len(tail)}章全是低张力",
                    "suggestion": "推进冲突或释放爽点"})
        return issues

    def _check_vocabulary_diversity(self, summaries: str, ch: int) -> List[Dict]:
        """[Bug6修复] 基于台账摘要而非完整正文"""
        if not summaries:
            return []
        # 提取最近20章摘要
        recent_text = ""
        for c in range(max(1, ch - 20), ch):
            m = re.search(rf'第{c}章[^\n]*\n(.*?)(?=\n第\d+章|\n## |\Z)', summaries, re.DOTALL)
            if m:
                recent_text += m.group(1)
        if len(recent_text) < 200:
            return []
        words = re.findall(r'[\u4e00-\u9fff]{2,4}', recent_text)
        if not words:
            return []
        ratio = len(set(words)) / len(words)
        if ratio < 0.15:
            return [{"severity": "info", "category": "词汇多样性",
                "description": f"近20章词汇多样性{ratio:.1%}",
                "suggestion": "增加同义词替换"}]
        return []

    def _check_sentence_patterns(self, summaries: str) -> List[Dict]:
        """[Bug6修复] 基于台账摘要"""
        if not summaries or len(summaries) < 200:
            return []
        sentences = [s.strip() for s in re.split(r'[。！？]', summaries) if len(s.strip()) > 2]
        if len(sentences) < 20:
            return []
        starters = [s[:2] for s in sentences]
        counter = Counter(starters)
        issues = []
        for starter, count in counter.most_common(3):
            ratio = count / len(sentences)
            if ratio > 0.08:
                issues.append({"severity": "info", "category": "句式分布",
                    "description": f"'{starter}'开头占{ratio:.1%}",
                    "suggestion": "变换句式开头"})
        return issues

    def _check_character_balance(self, summaries: str, ch: int) -> List[Dict]:
        if not summaries:
            return []
        last_seen: Dict[str, int] = {}
        for line in summaries.split("\n"):
            m = re.search(r'第(\d+)章', line)
            if m:
                c = int(m.group(1))
                names_m = re.search(r'出场人物[：:]\s*(.+?)(?:\s*\||$)', line)
                if names_m:
                    for name in re.split(r'[,，、]', names_m.group(1)):
                        name = name.strip()
                        if len(name) >= 2:
                            last_seen[name] = max(last_seen.get(name, 0), c)
        issues = []
        for name, last in last_seen.items():
            gap = ch - last
            if gap >= 20 and name not in ("旁白", "叙述"):
                issues.append({"severity": "info", "category": "角色出场均衡",
                    "description": f"'{name}'已{gap}章未出场",
                    "suggestion": "确认是否被遗忘"})
        return issues

    def _check_numerical_consistency(self, ledger: str) -> List[Dict]:
        if not ledger:
            return []
        issues = []
        for m in re.finditer(r'(\S{2,8}?)\s*[:：]\s*(-\d+)', ledger):
            issues.append({"severity": "warning", "category": "数值体系",
                "description": f"'{m.group(1)}'出现负数：{m.group(2)}",
                "suggestion": "检查扣减逻辑"})
        return issues


def format_audit_report(issues: List[Dict], ch: int) -> str:
    if not issues:
        return f"  [审计] 第{ch}章：全部通过"
    warnings = [i for i in issues if i["severity"] == "warning"]
    infos = [i for i in issues if i["severity"] == "info"]
    lines = [f"  [审计] 第{ch}章：{len(warnings)}个警告，{len(infos)}个提示"]
    for i in warnings:
        lines.append(f"    ⚠ [{i['category']}] {i['description']}")
        lines.append(f"      → {i['suggestion']}")
    for i in infos[:5]:
        lines.append(f"    ℹ [{i['category']}] {i['description']}")
    if len(infos) > 5:
        lines.append(f"    ℹ ...还有{len(infos) - 5}个提示")
    return "\n".join(lines)


# ============================================================
# DocxFormatter
# ============================================================

class DocxFormatter:

    def __init__(self, cfg: dict = None):
        c = cfg or CONFIG["docx"]
        self.font_body = c["font_body"]
        self.font_heading = c["font_heading"]
        self.size_body = Pt(c["size_body"])
        self.size_heading = Pt(c["size_heading"])
        self.line_spacing = c["line_spacing"]
        self.first_line_indent = Pt(c["first_line_indent"])
        self.margin = {k: Cm(c[f"margin_{k}"]) for k in ["top", "bottom", "left", "right"]}

    def create_chapter_docx(self, chapter_num: int, content: str, output_path: str):
        doc = Document()
        self._setup_page(doc)
        self._setup_styles(doc)
        self._add_title(doc, chapter_num)
        self._add_body(doc, content)
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    def _setup_page(self, doc):
        for s in doc.sections:
            s.page_width, s.page_height = Cm(21.0), Cm(29.7)
            for k, v in self.margin.items():
                setattr(s, f"{k}_margin", v)

    def _setup_styles(self, doc):
        style = doc.styles["Normal"]
        style.font.name = self.font_body
        style.font.size = self.size_body
        style.paragraph_format.line_spacing = self.line_spacing
        style.paragraph_format.first_line_indent = self.first_line_indent
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), self.font_body)

    def _add_title(self, doc, chapter_num):
        h = doc.add_heading(level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.paragraph_format.first_line_indent = Pt(0)
        run = h.add_run(f"第{chapter_num}章")
        run.font.name = self.font_heading
        run.font.size = self.size_heading
        run.font.bold = True
        run.element.rPr.rFonts.set(qn("w:eastAsia"), self.font_heading)

    def _add_body(self, doc, content):
        for line in content.strip().split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = self.font_body
                run.font.size = self.size_body
                run.element.rPr.rFonts.set(qn("w:eastAsia"), self.font_body)


# ============================================================
# ContextBuilder
# ============================================================

class ContextBuilder:

    def __init__(self, config_dir: str, memory_dir: str):
        self.config_dir = Path(config_dir)
        self.memory_dir = Path(memory_dir)
        self._story_bible_full: Optional[str] = None
        self._story_bible_sections: Dict[str, str] = {}
        self._chapter_outlines: Dict[int, str] = {}
        self._writing_rules: Optional[str] = None
        self._author_intent: Optional[str] = None
        self._current_focus: Optional[str] = None
        self._loaded = False

    def load_session_context(self) -> dict:
        stats = {}
        for name, attr, parser in [
            ("story_bible.md", "_story_bible_full", "bible"),
            ("chapter_outline.md", "_chapter_outlines", "outlines"),
            ("writing_rules.md", "_writing_rules", "text"),
            ("author_intent.md", "_author_intent", "text"),
            ("current_focus.md", "_current_focus", "text"),
        ]:
            path = self.config_dir / name
            if path.exists():
                raw = path.read_text(encoding="utf-8")
                if parser == "bible":
                    self._story_bible_full = raw
                    self._story_bible_sections = self._split_sections(raw)
                    stats["story_bible_chars"] = len(raw)
                elif parser == "outlines":
                    self._chapter_outlines = self._parse_chapter_outlines(raw)
                    stats["total_chapters_in_outline"] = len(self._chapter_outlines)
                elif parser == "text":
                    setattr(self, attr, raw)
                    stats[f"{name}_chars"] = len(raw)
        self._loaded = True
        return stats

    def build_main_context(self, chapter_num: int, memory: Dict[str, str], look_ahead: int = 5) -> Dict[str, str]:
        ctx = self.build_chapter_context(chapter_num, memory)
        ctx["upcoming_outlines"] = {
            i: self._chapter_outlines[i]
            for i in range(chapter_num + 1, chapter_num + look_ahead + 1)
            if i in self._chapter_outlines
        }
        ctx["story_bible_summary"] = self._story_bible_full[:1500] if self._story_bible_full else ""
        ctx["author_intent"] = self._author_intent or ""
        ctx["current_focus"] = self._current_focus or ""
        ctx["character_matrix"] = self._filter_relevant(memory.get("character_matrix", ""), ctx["chapter_outline"], 10)
        ctx["stale_subplots"] = self._detect_stale_subplots(memory.get("plot_threads", ""), chapter_num)
        return ctx

    def build_chapter_context(self, chapter_num: int, memory: Dict[str, str]) -> Dict[str, str]:
        if not self._loaded:
            raise RuntimeError("请先调用 load_session_context()")
        outline = self._chapter_outlines.get(chapter_num, f"第{chapter_num}章大纲未找到")
        return {
            "chapter_outline": outline,
            "writing_rules": self._writing_rules or "",
            "previous_summary": self._find_summary(chapter_num - 1, memory.get("chapter_summaries", "")),
            "character_states": self._filter_relevant(memory.get("character_tracking", ""), outline, 8),
            "character_matrix": self._filter_relevant(memory.get("character_matrix", ""), outline, 8),
            "active_plot_threads": self._filter_status(memory.get("plot_threads", ""), ["进行中", "活跃", "未完成", "待解决"], 5),
            "foreshadowing": self._filter_relevant(memory.get("foreshadowing", ""), outline, 5),
            "world_info": self._extract_world_info(outline, 800),
            "resource_ledger": self._filter_relevant(memory.get("resource_ledger", ""), outline, 5),
            "emotional_arcs": self._filter_relevant(memory.get("emotional_arcs", ""), outline, 5),
        }

    def _parse_chapter_outlines(self, text: str) -> Dict[int, str]:
        chapters, cur_num, cur_lines = {}, None, []
        for line in text.split('\n'):
            m = re.match(r'##[#]?\s+(?:第)?(\d+)\s*[章.:\-\s]', line)
            if m:
                if cur_num is not None:
                    chapters[cur_num] = '\n'.join(cur_lines).strip()
                cur_num, cur_lines = int(m.group(1)), [line]
            elif cur_num is not None:
                cur_lines.append(line)
        if cur_num is not None:
            chapters[cur_num] = '\n'.join(cur_lines).strip()
        return chapters

    def _split_sections(self, text: str) -> Dict[str, str]:
        sections = {}
        for part in re.split(r'\n(?=## )', text):
            m = re.match(r'##\s+(.+)', part)
            if m:
                sections[m.group(1).strip()] = part[m.end():].strip()
        return sections

    def _find_summary(self, num: int, text: str) -> str:
        if not text or num < 1:
            return "（无上一章摘要）"
        for pat in [rf'第{num}章[^\n]*\n(.*?)(?=\n第\d+章|\n## |\Z)',
                     rf'{num}\.\s*[^\n]*\n(.*?)(?=\n\d+\.\s|\n## |\Z)']:
            m = re.search(pat, text, re.DOTALL)
            if m:
                return m.group(1).strip()[:500]
        return f"（未找到第{num}章摘要）"

    def _filter_relevant(self, text: str, reference: str, max_items: int) -> str:
        if not text:
            return "（暂无记录）"
        keywords = self._keywords(reference)
        lines = [l.strip() for l in text.strip().split("\n") if l.strip()]
        if not keywords:
            return "\n".join(lines[-max_items:])
        scored = [(sum(1 for kw in keywords if kw in l), l) for l in lines]
        relevant = [(s, l) for s, l in scored if s > 0]
        if relevant:
            relevant.sort(key=lambda x: x[0], reverse=True)
            return "\n".join(l for _, l in relevant[:max_items])
        return "\n".join(lines[-max_items:])

    def _filter_status(self, text: str, keywords: List[str], max_items: int) -> str:
        if not text:
            return "（暂无记录）"
        lines = [l.strip() for l in text.strip().split("\n") if l.strip()]
        active = [l for l in lines if any(k in l for k in keywords)]
        return "\n".join(active[:max_items]) if active else "\n".join(lines[-max_items:])

    def _detect_stale_subplots(self, plot_text: str, ch: int) -> str:
        if not plot_text:
            return ""
        threshold = CONFIG["stale_subplot_threshold"]
        stale = []
        for line in plot_text.strip().split("\n"):
            line = line.strip()
            if not line or "已解决" in line or "已完成" in line:
                continue
            ch_matches = re.findall(r'第(\d+)章', line)
            if ch_matches:
                gap = ch - max(int(c) for c in ch_matches)
                if gap >= threshold:
                    stale.append(f"⚠ [{gap}章未推进] {line}")
        return "\n".join(stale) if stale else ""

    def _extract_world_info(self, outline: str, max_chars: int) -> str:
        if not self._story_bible_sections or not outline:
            return "（无特定世界设定需要强调）"
        keywords = self._keywords(outline)
        if not keywords:
            return "（无特定世界设定需要强调）"
        scored = []
        for title, content in self._story_bible_sections.items():
            score = sum(2 for kw in keywords if kw in title) + sum(1 for kw in keywords if kw in content[:300])
            if score > 0:
                scored.append((score, f"【{title}】\n{content}"))
        scored.sort(key=lambda x: x[0], reverse=True)
        result = ""
        for _, section in scored:
            if len(result) + len(section) > max_chars:
                remaining = max_chars - len(result)
                if remaining > 50:
                    result += section[:remaining] + "..."
                break
            result += section + "\n\n"
        return result.strip() if result else "（无特定世界设定需要强调）"

    def _keywords(self, text: str) -> List[str]:
        cn = re.findall(r'[\u4e00-\u9fff]{2,4}', text)
        en = [w for w in re.findall(r'[A-Za-z]{3,}', text)]
        stop = {"一个", "这个", "那个", "他们", "她们", "我们", "你们", "什么", "怎么", "为什么",
                "可以", "没有", "但是", "然后", "因为", "所以", "如果", "虽然", "只是", "已经",
                "正在", "开始", "发现", "决定", "告诉", "知道", "认为", "觉得", "来到", "看到",
                "听到", "想到", "说道", "chapter"}
        words = [w for w in cn + en if w.lower() not in stop]
        seen = set()
        return [w for w in words if not (w in seen or seen.add(w))][:30]


# ============================================================
# CheckpointManager
# ============================================================

class CheckpointManager:

    def __init__(self, checkpoint_path: str, output_dir: str, total_chapters: int):
        self.path = Path(checkpoint_path)
        self.output_dir = Path(output_dir)
        self.total_chapters = total_chapters
        self.state = self._load_or_init()

    def _load_or_init(self) -> dict:
        if self.path.exists():
            return json.loads(self.path.read_text(encoding="utf-8"))
        return {"total_chapters": self.total_chapters, "completed": [], "last_completed": 0,
                "memory_synced_to": 0, "created": datetime.now().isoformat(), "updated": datetime.now().isoformat()}

    def scan_existing_chapters(self) -> List[int]:
        if not self.output_dir.exists():
            return []
        chapters = []
        for f in sorted(self.output_dir.iterdir()):
            if f.suffix in ('.md', '.txt', '.docx'):
                m = re.search(r'(?:chapter_|第)(\d+)', f.stem)
                if m:
                    chapters.append(int(m.group(1)))
        return sorted(set(chapters))

    def sync_with_files(self) -> dict:
        existing = self.scan_existing_chapters()
        cp_set, fs_set = set(self.state["completed"]), set(existing)
        report = {"checkpoint_says": len(cp_set), "files_found": len(fs_set),
                  "in_checkpoint_not_files": sorted(cp_set - fs_set), "in_files_not_checkpoint": sorted(fs_set - cp_set)}
        self.state["completed"] = sorted(fs_set)
        self.state["last_completed"] = max(fs_set) if fs_set else 0
        self.save()
        return report

    def get_resume_point(self) -> int:
        return self.state["last_completed"] + 1 if self.state["completed"] else 1

    def mark_complete(self, chapter_num: int):
        if chapter_num not in self.state["completed"]:
            self.state["completed"].append(chapter_num)
            self.state["completed"].sort()
        self.state["last_completed"] = max(self.state["completed"])
        self.state["updated"] = datetime.now().isoformat()
        self.save()

    def mark_memory_synced(self, chapter_num: int):
        self.state["memory_synced_to"] = max(self.state["memory_synced_to"], chapter_num)
        self.save()

    def get_memory_gap(self, up_to: int) -> List[int]:
        synced = self.state["memory_synced_to"]
        completed = set(self.state["completed"])
        return [ch for ch in range(synced + 1, up_to + 1) if ch in completed]

    def save(self):
        self.path.write_text(json.dumps(self.state, ensure_ascii=False, indent=2), encoding="utf-8")

    def get_status(self) -> str:
        c = len(self.state["completed"])
        t = self.state["total_chapters"]
        return f"已完成: {c}/{t} 章 | 最新: 第{self.state['last_completed']}章 | 记忆同步到: 第{self.state['memory_synced_to']}章"


# ============================================================
# MemorySystem
# ============================================================

class MemorySystem:
    FILES = [
        "global_memory.md", "character_tracking.md", "character_matrix.md",
        "plot_threads.md", "foreshadowing.md", "chapter_summaries.md",
        "resource_ledger.md", "emotional_arcs.md", "world_info.md", "timeline.md",
    ]

    def __init__(self, memory_dir: str):
        self.dir = Path(memory_dir)
        self.dir.mkdir(parents=True, exist_ok=True)
        self._data: Dict[str, str] = {}
        self.reload()

    def reload(self):
        for fname in self.FILES:
            key = fname.replace(".md", "")
            path = self.dir / fname
            self._data[key] = path.read_text(encoding="utf-8") if path.exists() else ""

    def get_all(self) -> Dict[str, str]:
        return dict(self._data)

    def update_file(self, key: str, content: str):
        path = self.dir / f"{key}.md"
        atomic_write(path, content)
        self._data[key] = content


# ============================================================
# MemoryUpdater [Bug3修复: 崩溃保护]
# ============================================================

class MemoryUpdater:
    MAPPING = {
        "角色变化": "character_tracking",
        "角色信息边界": "character_matrix",
        "情节推进": "plot_threads",
        "伏笔动态": "foreshadowing",
        "章节摘要": "chapter_summaries",
        "资源变动": "resource_ledger",
        "情感变化": "emotional_arcs",
        "世界设定更新": "world_info",
    }

    def __init__(self, memory_dir: str):
        self.dir = Path(memory_dir)
        self.dir.mkdir(parents=True, exist_ok=True)

    async def update(self, chapter_num: int, content: str):
        prompt = f"""根据第{chapter_num}章正文，列出以下信息（纯文本，每行一条，用"|"分隔字段）：

角色变化：角色名 | 当前状态 | 变化说明
角色信息边界：角色名 | 本章获知的信息 | 本章仍不知道的信息
情节推进：情节线名称 | 状态（进行中/已解决/新出现） | 最近活跃: 第{chapter_num}章 | 简述
伏笔动态：伏笔内容 | 状态（已埋设/已回收/进行中） | 简述
章节摘要：（100字以内的本章摘要）
资源变动：物品名 | 归属/状态变化 | 备注
情感变化：角色名 | 情绪变化 | 触发事件
世界设定更新：（如有新设定，列出；无则写"无"）

## 第{chapter_num}章正文
{content}"""
        result = await call_llm_api(prompt, temperature=0.3)
        parsed = self._safe_parse(result)
        self._apply_safely(chapter_num, parsed)

    def apply_from_meta(self, chapter_num: int, meta: str, summary: str):
        sections = {label: [] for label in self.MAPPING}
        current_key = None
        for line in meta.split("\n"):
            line = line.strip()
            if not line:
                continue
            matched = False
            for label in self.MAPPING:
                if label in line:
                    current_key = label
                    matched = True
                    after = line.split(label, 1)[-1].strip("：: ")
                    if after and after != "无":
                        sections[label].append(after)
                    break
            if not matched and current_key:
                if line != "无":
                    sections[current_key].append(line)
        if summary:
            sections["章节摘要"] = [summary]
        self._apply_safely(chapter_num, sections)

    def _safe_parse(self, text: str) -> dict:
        result = {label: [] for label in self.MAPPING}
        result["章节摘要"] = []
        current_key = None
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            matched = False
            for label in self.MAPPING:
                if label in line:
                    current_key = label
                    matched = True
                    after = line.split(label, 1)[-1].strip("：: ")
                    if after and after != "无":
                        result[label].append(after)
                    break
            if not matched and current_key:
                if line != "无":
                    result[current_key].append(line)
        return result

    def _apply_safely(self, chapter_num: int, parsed: dict):
        """[崩溃保护] 使用 atomic_write 防止写入一半崩溃"""
        for label, file_key in self.MAPPING.items():
            try:
                items = parsed.get(label, [])
                if not items:
                    continue
                path = self.dir / f"{file_key}.md"
                existing = path.read_text(encoding="utf-8") if path.exists() else ""
                if label == "章节摘要":
                    entry = f"\n\n## 第{chapter_num}章\n" + "\n".join(items)
                else:
                    entry = f"\n\n## 第{chapter_num}章更新\n" + "\n".join(f"- {item}" for item in items)
                atomic_write(path, existing + entry)
            except Exception as e:
                print(f"  [警告] {file_key} 更新失败: {e}")


# ============================================================
# SubprocessWriter
# ============================================================

class SubprocessWriter:

    async def write_chapter(self, chapter_num: int, context: Dict[str, str], decision: str, word_count_target: int = 3000) -> str:
        prompt = f"""你是一位专业的网络小说写手。请根据以下信息写作第{chapter_num}章。

## 本章大纲
{context.get('chapter_outline', '无')}

## 主进程写作指示
{decision}

## 写作规则
{context.get('writing_rules', '')}

## 上一章摘要（用于衔接）
{context.get('previous_summary', '这是第一章')}

## 相关角色当前状态
{context.get('character_states', '暂无')}

## 角色信息边界（严格遵守，禁止信息越界）
{context.get('character_matrix', '暂无')}

## 活跃情节线
{context.get('active_plot_threads', '暂无')}

## 需要回收的伏笔
{context.get('foreshadowing', '暂无')}

## 相关世界设定
{context.get('world_info', '无特定设定')}

## 资源账本
{context.get('resource_ledger', '暂无')}

## 情感弧线
{context.get('emotional_arcs', '暂无')}

## 写作要求
1. 目标字数：{word_count_target}字左右
2. 保持与上一章的自然衔接
3. 体现主进程指示中的重点
4. 章末留悬念/钩子
5. 【信息边界】角色只能基于已知信息行动和思考

## 输出格式（严格遵守）

=== CHAPTER_TITLE ===
（章节标题，不含"第X章"）

=== CHAPTER_CONTENT ===
（正文内容，目标{word_count_target}字）

=== CHAPTER_SUMMARY ===
（100字以内摘要：出场人物、关键事件、状态变化、伏笔动态）

=== META ===
角色变化：角色名 | 当前状态 | 变化说明
角色信息边界：角色名 | 本章获知的信息 | 本章仍不知道的信息
情节推进：情节线名称 | 状态 | 最近活跃: 第{chapter_num}章 | 简述
伏笔动态：伏笔内容 | 状态 | 简述
资源变动：物品名 | 归属/状态变化
情感变化：角色名 | 情绪变化

请开始写作第{chapter_num}章："""
        return await call_llm_api(prompt, temperature=0.7, max_tokens=8000)

    async def revise_chapter(self, chapter_num: int, content: str, issues: list) -> str:
        issue_text = "\n".join(f"- [{i['category']}] {i['suggestion']}" for i in issues)
        prompt = f"""请修订第{chapter_num}章正文，消除以下问题：

{issue_text}

## 原文
{content}

## 修订要求
1. 只修改有问题的部分
2. 保持情节、人物、对话不变
3. 保持 === CHAPTER_TITLE === === CHAPTER_CONTENT === === CHAPTER_SUMMARY === === META === 格式"""
        return await call_llm_api(prompt, temperature=0.5)


# ============================================================
# LLM API
# ============================================================

async def call_llm_api(prompt: str, temperature: float = 0.7, max_tokens: int = 8000) -> str:
    raise NotImplementedError("请实现 call_llm_api 函数，对接你的 LLM API")


# ============================================================
# 输出解析器
# ============================================================

def parse_chapter_output(raw: str) -> Dict[str, str]:
    result = {"title": "", "content": "", "summary": "", "meta": ""}
    key_map = {"CHAPTER_TITLE": "title", "CHAPTER_CONTENT": "content", "CHAPTER_SUMMARY": "summary", "META": "meta"}
    current_key = None
    current_lines = []
    for line in raw.split("\n"):
        marker = re.match(r'^===\s*(\w+)\s*===\s*$', line.strip())
        if marker:
            tag = marker.group(1)
            if tag in key_map:
                if current_key:
                    result[current_key] = "\n".join(current_lines).strip()
                current_key = key_map[tag]
                current_lines = []
                continue
        if current_key is not None:
            current_lines.append(line)
    if current_key:
        result[current_key] = "\n".join(current_lines).strip()
    if not result["content"]:
        result["content"] = raw
    return result


# ============================================================
# NovelOrchestrator - 主进程
# ============================================================

class NovelOrchestrator:

    def __init__(self, total_chapters: int = 500):
        self.config_dir = CONFIG["config_dir"]
        self.memory_dir = CONFIG["memory_dir"]
        self.output_dir = CONFIG["output_dir"]

        self.context_builder = ContextBuilder(self.config_dir, self.memory_dir)
        self.checkpoint = CheckpointManager(CONFIG["checkpoint_path"], self.output_dir, total_chapters)
        self.memory = MemorySystem(self.memory_dir)
        self.writer = SubprocessWriter()
        self.memory_updater = MemoryUpdater(self.memory_dir)
        self.docx_formatter = DocxFormatter()
        self.validator = PostWriteValidator()
        self.audit = AuditSystem(self.memory_dir, self.output_dir)

        Path(self.output_dir).mkdir(parents=True, exist_ok=True)

    async def run(self, start_chapter: Optional[int] = None, end_chapter: Optional[int] = None, rebuild_memory: bool = False):
        print("=" * 60)
        print("NovelForge v2.2.2")
        print("=" * 60)

        print("\n正在加载会话上下文...")
        stats = self.context_builder.load_session_context()
        for k, v in stats.items():
            print(f"  {k}: {v}")

        print("\n正在同步checkpoint...")
        report = self.checkpoint.sync_with_files()
        print(f"  checkpoint: {report['checkpoint_says']} 章 | 文件: {report['files_found']} 章")
        if report["in_files_not_checkpoint"]:
            print(f"  新发现: {report['in_files_not_checkpoint']}")
        print(f"\n{self.checkpoint.get_status()}")

        if start_chapter is None:
            start_chapter = self.checkpoint.get_resume_point()
        if end_chapter is None:
            end_chapter = self.checkpoint.state["total_chapters"]

        print(f"\n写作范围: 第{start_chapter}章 → 第{end_chapter}章 ({end_chapter - start_chapter + 1}章)")

        if rebuild_memory:
            await self._rebuild_memory(start_chapter - 1)

        gap = self.checkpoint.get_memory_gap(start_chapter - 1)
        if gap:
            print(f"\n⚠ 记忆缺口: 第{gap[0]}~{gap[-1]}章")
            await self._rebuild_memory_for_chapters(gap)

        print(f"\n{'=' * 60}\n准备就绪\n{'=' * 60}")

        for chapter_num in range(start_chapter, end_chapter + 1):
            await self._write_one_chapter(chapter_num)

        print(f"\n{'=' * 60}\n全部完成！{self.checkpoint.get_status()}")

    async def _write_one_chapter(self, chapter_num: int):
        print(f"\n{'─' * 40}")
        print(f"第{chapter_num}章")
        print(f"{'─' * 40}")

        memory_data = self.memory.get_all()

        # ---- 主进程决策 ----
        main_context = self.context_builder.build_main_context(chapter_num, memory_data, look_ahead=5)
        if main_context.get("stale_subplots"):
            print(f"  ⚠ 支线停滞:")
            for line in main_context["stale_subplots"].split("\n"):
                if line.strip():
                    print(f"    {line.strip()}")

        print(f"  [主进程] 分析上下文...")
        decision = await self._make_decision(chapter_num, main_context)

        # ---- 子进程写作 ----
        sub_context = self.context_builder.build_chapter_context(chapter_num, memory_data)
        print(f"  [子进程] 写作中...")
        raw_output = await self.writer.write_chapter(chapter_num, sub_context, decision, CONFIG["word_count_target"])

        # ---- 解析输出 ----
        parsed = parse_chapter_output(raw_output)
        content = parsed["content"]

        # ---- 写后验证：硬规则 spot-fix ----
        content, fixes = self.validator.spot_fix(content)
        if fixes:
            print(f"  [Spot-fix] {', '.join(fixes)}")

        # ---- AI 痕迹检测 + 修订循环 ----
        ai_issues = analyze_ai_tells(content)
        ai_warnings = [i for i in ai_issues if i["severity"] == "warning"]

        if ai_warnings:
            print(f"  [AI检测] {format_ai_tell_report(ai_issues)}")
            content = await self._ai_revise_loop(chapter_num, content, ai_warnings, parsed)
        else:
            print(f"  [AI检测] 通过")

        # ---- 18 维度审计 ----
        memory_data = self.memory.get_all()  # 重新读取（可能因 spot-fix 更新）
        audit_issues = self.audit.run_audit(chapter_num, content, memory_data)

        if audit_issues:
            print(format_audit_report(audit_issues, chapter_num))

            # [Gap1修复] 审计→修订循环
            fixable = self.audit.get_fixable_warnings(audit_issues)
            if fixable:
                content = await self._audit_revise_loop(chapter_num, content, fixable, parsed)
        else:
            print(f"  [审计] 通过")

        # ---- 保存 ----
        self._save_chapter(chapter_num, content)
        print(f"  [保存] chapter_{chapter_num:03d}.md + .docx")

        # ---- 更新记忆 ----
        print(f"  [记忆] 更新台账...")
        if parsed["meta"]:
            self.memory_updater.apply_from_meta(chapter_num, parsed["meta"], parsed["summary"])
        else:
            await self.memory_updater.update(chapter_num, content)
        self.memory.reload()

        # ---- checkpoint ----
        self.checkpoint.mark_complete(chapter_num)
        self.checkpoint.mark_memory_synced(chapter_num)
        print(f"  [完成] 第{chapter_num}章")

    async def _ai_revise_loop(self, chapter_num: int, content: str, warnings: list, parsed: dict) -> str:
        """AI痕迹修订循环"""
        for attempt in range(CONFIG["max_rewrite_attempts"]):
            print(f"  [AI修订] 第{attempt + 1}次...")
            revised = await self.writer.revise_chapter(chapter_num, content, warnings)
            revised_parsed = parse_chapter_output(revised)
            revised_content = revised_parsed["content"]
            revised_content, _ = self.validator.spot_fix(revised_content)

            re_issues = analyze_ai_tells(revised_content)
            re_warnings = [i for i in re_issues if i["severity"] == "warning"]

            if len(re_warnings) < len(warnings):
                content = revised_content
                if revised_parsed["summary"]:
                    parsed["summary"] = revised_parsed["summary"]
                if revised_parsed["meta"]:
                    parsed["meta"] = revised_parsed["meta"]
                warnings = re_warnings
                if not re_warnings:
                    print(f"  [AI修订] 已清除")
                    break
            else:
                break
        return content

    async def _audit_revise_loop(self, chapter_num: int, content: str, fixable: list, parsed: dict) -> str:
        """[Gap1修复] 审计→修订循环"""
        for attempt in range(CONFIG["max_audit_revise_attempts"]):
            print(f"  [审计修订] 第{attempt + 1}次（{len(fixable)}个问题）...")
            revised = await self.writer.revise_chapter(chapter_num, content, fixable)
            revised_parsed = parse_chapter_output(revised)
            revised_content = revised_parsed["content"]
            revised_content, _ = self.validator.spot_fix(revised_content)

            # 重新审计
            memory_data = self.memory.get_all()
            re_issues = self.audit.run_audit(chapter_num, revised_content, memory_data)
            re_fixable = self.audit.get_fixable_warnings(re_issues)

            if len(re_fixable) < len(fixable):
                content = revised_content
                if revised_parsed["summary"]:
                    parsed["summary"] = revised_parsed["summary"]
                if revised_parsed["meta"]:
                    parsed["meta"] = revised_parsed["meta"]
                fixable = re_fixable
                if not re_fixable:
                    print(f"  [审计修订] 已修复")
                    break
            else:
                print(f"  [审计修订] 无改善，保留当前版本")
                break
        return content

    async def _make_decision(self, chapter_num: int, ctx: dict) -> str:
        upcoming_text = self._format_upcoming(ctx.get("upcoming_outlines", {}))
        stale_text = ctx.get("stale_subplots", "")

        prompt = f"""你是一位网络小说的总编辑。请为第{chapter_num}章制定写作指示。

## 本章大纲
{ctx['chapter_outline']}

## 后续章节大纲
{upcoming_text}

## 上一章摘要
{ctx['previous_summary']}

## 相关角色状态
{ctx['character_states']}

## 角色信息边界
{ctx.get('character_matrix', '暂无')}

## 活跃情节线
{ctx['active_plot_threads']}

## 待回收伏笔
{ctx['foreshadowing']}

## 资源账本
{ctx.get('resource_ledger', '暂无')}

## 情感弧线
{ctx.get('emotional_arcs', '暂无')}

## 相关世界设定
{ctx['world_info']}

## 故事圣经摘要
{ctx.get('story_bible_summary', '')}

## 作者长期意图
{ctx.get('author_intent', '')}

## 当前关注点
{ctx.get('current_focus', '')}

{"## ⚠ 支线停滞警告" + chr(10) + stale_text if stale_text else ""}

请输出（不超过500字）：
1. 本章重点
2. 节奏建议
3. 角色互动要点
4. 信息边界提醒
5. 连续性细节
6. 章末钩子建议
7. 本章禁忌"""
        return await call_llm_api(prompt, temperature=0.5)

    async def _rebuild_memory(self, up_to: int):
        chapters = [ch for ch in self.checkpoint.state["completed"] if ch <= up_to]
        if chapters:
            await self._rebuild_memory_for_chapters(chapters)

    async def _rebuild_memory_for_chapters(self, chapters: list):
        total = len(chapters)
        print(f"重建记忆: {total} 章")
        for i, ch in enumerate(chapters, 1):
            path = Path(self.output_dir) / f"chapter_{ch:03d}.md"
            if not path.exists():
                alt = list(Path(self.output_dir).glob(f"*{ch:03d}*"))
                path = alt[0] if alt else None
            if not path or not path.exists():
                print(f"  [{i}/{total}] 跳过第{ch}章")
                continue
            content = path.read_text(encoding="utf-8")
            await self.memory_updater.update(ch, content)
            self.checkpoint.mark_memory_synced(ch)
            print(f"  [{i}/{total}] 第{ch}章")
        self.memory.reload()
        print(f"记忆重建完成")

    def _save_chapter(self, chapter_num: int, content: str):
        md_path = Path(self.output_dir) / f"chapter_{chapter_num:03d}.md"
        md_path.write_text(content, encoding="utf-8")
        docx_path = Path(self.output_dir) / f"chapter_{chapter_num:03d}.docx"
        self.docx_formatter.create_chapter_docx(chapter_num, content, str(docx_path))

    def _format_upcoming(self, upcoming: dict) -> str:
        if not upcoming:
            return "（无后续大纲）"
        return "\n".join(f"第{n}章: {o[:200]}{'...' if len(o) > 200 else ''}" for n, o in sorted(upcoming.items()))


# ============================================================
# 写作规则模板
# ============================================================

def generate_writing_rules_template(output_path: str = "config/writing_rules.md"):
    content = """# 写作规则 v2.2.2

## 核心规则

1. 以简体中文工作，句子长短交替，段落适合手机阅读（3-5行/段）
2. 伏笔前后呼应，不留悬空线
3. 只读必要上下文，不机械重复已有内容

## 人物塑造铁律

- 人设一致性：角色行为由"过往经历 + 当前利益 + 性格底色"共同驱动
- 人物立体化：核心标签 + 反差细节 = 活人
- 拒绝工具人：配角必须有独立动机和反击能力
- 角色区分度：不同角色的语气、发怒方式、处事模式必须有显著差异
- 情感/动机逻辑链：任何关系改变都必须有铺垫和事件驱动

## 信息边界铁律

- 【铁律】每个角色只能基于已知信息行动和思考
- 【铁律】角色A没有亲历或被告知事件B，A不能对B做出反应
- 【铁律】角色对局势的误判必须基于其有限的信息边界
- 【铁律】多角色同场景时，每个角色的台词只能包含该角色已知的信息
- 【铁律】信息传递必须有明确的传递链

## 叙事技法

- Show, don't tell：用细节堆砌真实，用行动证明强大
- 五感代入法：场景描写中加入1-2种五感细节
- 钩子设计：每章结尾设置悬念/伏笔/钩子
- 对话驱动：优先用对话传递冲突和信息
- 信息分层植入：严禁大段灌输世界观
- 描写必须服务叙事
- 日常/过渡段落必须为后续剧情服务

## 逻辑自洽

- 三连反问自检："为什么这么做？""符合利益吗？""符合人设吗？"
- 信息越界检查
- 关系改变必须事件驱动
- 场景转换必须有过渡
- 每段至少带来一项新信息

## 语言约束

- 句式多样化：严禁连续相同句式或相同主语开头
- 多用动词和名词，少用形容词
- 群像反应具体到个人
- 情绪用细节传达
- 禁止元叙事

## 去AI味铁律

- 【铁律】叙述者永远不得替读者下结论
- 【铁律】禁止分析报告式语言（核心动机、信息边界、信息落差、核心风险、利益最大化、当前处境）
- 【铁律】转折/惊讶标记词（仿佛、忽然、竟、竟然、猛地、猛然、不禁、宛如）全篇每3000字最多1次
- 【铁律】同一体感/意象禁止连续渲染超过两轮
- 【铁律】六步走心理分析术语只用于内部推理，不出现在正文
- 【硬性禁令】禁止"不是……而是……"句式
- 【硬性禁令】禁止破折号"——"
- 禁止账本式数据出现在正文

## 反例→正例对照

### 情绪描写
| 反例 | 正例 | 要点 |
|---|---|---|
| 他感到非常愤怒。 | 他捏碎了手中的茶杯，滚烫的茶水流过指缝。 | 用动作外化情绪 |
| 她心里很悲伤。 | 她攥紧手机，指节发白，屏幕上的聊天记录模糊成一片。 | 身体细节替代标签 |
| 他感到一阵恐惧。 | 他后背的汗毛竖了起来，脚底像踩在了冰上。 | 五感传递 |

### 转折与衔接
| 反例 | 正例 | 要点 |
|---|---|---|
| 虽然他很强，但是他还是输了。 | 他确实强，可对面那个老东西更脏。 | 口语化转折 |
| 然而，事情并没有那么简单。 | 哪有那么便宜的事。 | 角色内心吐槽 |
| 因此，他决定采取行动。 | 他站起来，把凳子踢到一边。 | 直接写动作 |

### 叙述者姿态
| 反例 | 正例 | 要点 |
|---|---|---|
| 这一刻，他终于明白了什么是力量。 | 删掉，让读者自己感受。 | 不替读者下结论 |
| 显然，对方低估了他。 | 只写对方的表情变化。 | "显然"是说教 |
| 全场为之震惊。 | 老陈的烟掉在裤子上，烫得他跳起来。 | 具体到个人 |

## 六步走人物心理分析

1. **当前处境**：角色面临什么局面？
2. **核心动机**：最想要什么？最害怕什么？
3. **信息边界**：知道什么？不知道什么？
4. **性格过滤**：性格会怎么反应？
5. **行为选择**：基于以上四点做出什么选择？
6. **情绪外化**：伴随什么情绪？

### 人设防崩三问
1. "为什么这么做？"——必须有驱动
2. "符合之前的人设吗？"
3. "读者会觉得突兀吗？"

### "盐溶于汤"原则
价值观通过行为传达，不喊口号。

## 配角设计方法论

- 配角必须有自己的算盘和反击
- 动机绑定主线
- 核心标签 + 反差细节
- 通过事件立人设
- 语言区分度
- 拒绝集体反应

## 读者心理学框架

- **期待管理**：适当延迟释放，增强快感
- **信息落差**：让读者比角色多知道或少知道一点
- **情绪节拍**：压制→释放→更大的压制→更大的释放
- **锚定效应**：先给参照，再展示表现
- **沉没成本**：每章给出继续读的理由
- **代入感维护**：困境让读者共情，选择让读者认同

## 情感节点设计

1. 设计3-5个关键事件
2. 递进升温，禁止跨越式发展
3. 情绪用场景传达
4. 情感与题材匹配
5. 禁止标签化互动

### 强情绪升级法
1. 加入前因后果
2. 坏事叠坏事
3. 日常必须为主线服务

## 代入感六支柱

1. 基础信息标签化：一百字内知道谁、在哪、什么
2. 可视化熟悉感：读者碰过的具体细节
3. 共鸣分两层：认知共鸣 + 情绪共鸣
4. 欲望两条腿：基础欲望 + 主动欲望
5. 五感钩子：除视觉外放1-2种感官
6. 人设活化：核心标签 + 反差细节

## 创作宪法

1. Show don't tell
2. 价值观像盐溶于汤
3. 行动立于三条腿：过往经历、当前利益、性格底色
4. 配角都有自己的账本
5. 节奏即呼吸
6. 每章结尾有钩子
7. 全员智商在线
8. 后世梗符合年代
9. 时间线不能错
10. 日常七成必须成为伏笔
11. 关系改变要事件驱动
12. 人设前后一致
13. 重要剧情用场景不用总结
14. 拒绝流水账

## 黄金三章纪律

### 第一章：抛出核心冲突
- 开篇直接进入冲突
- 第一段必须有动作或对话
- 最多1-2个场景，3个角色
- 身份通过行动带出
- 核心矛盾浮出水面

### 第二章：展现金手指
- 核心优势初现
- 通过具体事件展现
- 第一个小爽点

### 第三章：明确短期目标
- 具体可衡量的目标
- 章尾强钩子

## 支线管理

- 每条支线5章内至少推进一次
- 停滞支线必须决定：推进或给出理由
- 不能所有支线同时推进或同时停滞
- 支线必须在2-3章内与主线关联

## AI高频词黑名单

### 高疲劳词（单章最多1次）
仿佛、忽然、竟然、猛地、猛然、不禁、宛如、随即、旋即、霎时、顿时

### 套话词（每千字不超过3次）
似乎、可能、或许、大概、某种程度上、一定程度上

### 公式化转折（同一词单章不超过2次）
然而、不过、与此同时、另一方面、尽管如此、话虽如此

### 结构检测
- 段落长度变异系数不低于0.15
- 连续3句以上相同开头视为违规

## docx 输出格式

- 纸张：A4 | 页边距：上下2.54cm 左右3.18cm
- 正文：宋体 小四 1.5倍行距 首行缩进2字符
- 标题：黑体 小三 加粗 居中
"""
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text(content, encoding="utf-8")
    print(f"写作规则模板已生成: {output_path}")


# ============================================================
# 入口
# ============================================================

if __name__ == "__main__":
    import sys

    if "--generate-rules" in sys.argv:
        generate_writing_rules_template()
        sys.exit(0)

    total = int(sys.argv[1]) if len(sys.argv) > 1 else 500
    start = int(sys.argv[2]) if len(sys.argv) > 2 else None
    rebuild = "--rebuild" in sys.argv

    print(f"NovelForge v2.2.2 | 总章数: {total}")
    if start:
        print(f"起始章: {start}")
    if rebuild:
        print("模式: 重建记忆后续写")

    asyncio.run(NovelOrchestrator(total_chapters=total).run(
        start_chapter=start,
        rebuild_memory=rebuild,
    ))
