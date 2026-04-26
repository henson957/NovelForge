#!/usr/bin/env python3
"""
NovelForge v2.0.6
长篇小说全自动写作系统
Python + AI 混合架构

安装依赖：
pip install openai python-docx

使用方式：
python novelforge.py
"""

import os
import re
import json
import time
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass
from typing import Optional, Tuple

from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class OpenClawConfigReader:

    CONFIG_PATH = os.path.expanduser("~/.openclaw/openclaw.json")

    @staticmethod
    def read() -> dict:
        config_path = OpenClawConfigReader.CONFIG_PATH

        if not os.path.exists(config_path):
            print(f"[配置] 未找到 OpenClaw 配置文件：{config_path}")
            print("[配置] 将使用手动配置或默认值")
            return {}

        try:
            with open(config_path, "r", encoding="utf-8") as f:
                raw = json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            print(f"[配置] 读取 OpenClaw 配置失败：{e}")
            return {}

        result = {
            "api_key": "",
            "base_url": "",
            "model": "",
            "provider": "",
        }

        agent_config = raw.get("agents", {}).get("main", {})
        if not agent_config:
            for key in raw.get("agents", {}):
                agent_config = raw["agents"][key]
                break

        model_str = agent_config.get("model", "")
        if "/" in model_str:
            provider, model = model_str.split("/", 1)
            result["provider"] = provider
            result["model"] = model
        else:
            result["model"] = model_str
            result["provider"] = agent_config.get("provider", "")

        providers = raw.get("providers", {})
        provider_name = result["provider"]

        if provider_name and provider_name in providers:
            prov = providers[provider_name]
            result["api_key"] = prov.get("apiKey", "") or prov.get("api_key", "")
            result["base_url"] = prov.get("baseUrl", "") or prov.get("base_url", "")

        if not result["api_key"]:
            result["api_key"] = raw.get("apiKey", "") or raw.get("api_key", "")
        if not result["base_url"]:
            result["base_url"] = raw.get("baseUrl", "") or raw.get("base_url", "")

        if not result["api_key"] or not result["base_url"]:
            for node in raw.get("nodes", []):
                if node.get("provider", "") == provider_name:
                    if not result["api_key"]:
                        result["api_key"] = node.get("apiKey", "") or node.get("api_key", "")
                    if not result["base_url"]:
                        result["base_url"] = node.get("baseUrl", "") or node.get("base_url", "")
                    break

        return result

    @staticmethod
    def validate(config: dict) -> Tuple[bool, list]:
        missing = []
        if not config.get("api_key"):
            missing.append("api_key")
        if not config.get("base_url"):
            missing.append("base_url")
        if not config.get("model"):
            missing.append("model")
        return len(missing) == 0, missing


@dataclass
class Config:
    ai_base_url: str = ""
    ai_api_key: str = ""
    ai_model: str = ""
    ai_timeout: int = 300

    project_name: str = ""
    total_chapters: int = 0
    words_per_chapter: int = 2500

    review_standards: str = ""
    writing_style: str = ""
    special_requirements: str = ""

    chapter_retry_max: int = 3
    batch_review_interval: int = 50
    recent_summary_window: int = 8
    transition_word_interval: int = 2000
    final_review_max_rounds: int = 3

    base_path: str = ""
    original_path: str = ""
    memory_path: str = ""
    chapter_path: str = ""
    review_path: str = ""

    def auto_load_api_config(self) -> bool:
        oc_config = OpenClawConfigReader.read()

        valid, missing = OpenClawConfigReader.validate(oc_config)
        if valid:
            self.ai_api_key = oc_config["api_key"]
            self.ai_base_url = oc_config["base_url"]
            self.ai_model = oc_config["model"]
            print(f"[配置] 已从 OpenClaw 自动加载 API 配置")
            print(f"[配置] 模型：{oc_config['provider']}/{oc_config['model']}")
            return True
        else:
            print(f"[配置] OpenClaw 配置不完整，缺少：{', '.join(missing)}")
            print("[配置] 请手动提供 API 配置")
            return False

    def setup_paths(self):
        self.base_path = os.path.join(os.path.expanduser("~"), "Desktop", f"{self.project_name}项目")
        self.original_path = os.path.join(self.base_path, "原始材料")
        self.memory_path = os.path.join(self.base_path, "记忆系统")
        self.chapter_path = os.path.join(self.base_path, "章节")
        self.review_path = os.path.join(self.base_path, "审核记录")
        for p in [self.original_path, self.memory_path, self.chapter_path, self.review_path]:
            os.makedirs(p, exist_ok=True)


class AIInterface:

    def __init__(self, config: Config):
        self.client = OpenAI(
            base_url=config.ai_base_url,
            api_key=config.ai_api_key,
        )
        self.model = config.ai_model
        self.timeout = config.ai_timeout

    def generate(self, system_prompt: str, user_prompt: str) -> str:
        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            timeout=self.timeout,
        )
        return response.choices[0].message.content


class MemorySystem:

    def __init__(self, config: Config):
        self.config = config
        self.mp = config.memory_path

    def initialize(self, story_bible: str, chapter_outlines: str):
        outline_index = self._build_outline_index(chapter_outlines)
        self._write_file("全局记忆.docx", "（等待AI从故事圣经中提取）")
        self._write_file("线索台账.docx", "（等待AI从故事圣经中提取）")
        self._write_file("近期摘要.docx", "（初始为空，写完第一章后开始填充）")
        self._write_file("章节大纲索引.docx", outline_index)
        self._write_file("写作进度.docx", self._format_progress(0, 1, "写作中"))

    def _build_outline_index(self, chapter_outlines: str) -> str:
        return chapter_outlines

    def read_global_memory(self) -> str:
        return self._read_file("全局记忆.docx")

    def read_clue_ledger(self) -> str:
        return self._read_file("线索台账.docx")

    def read_recent_summaries(self) -> str:
        return self._read_file("近期摘要.docx")

    def read_chapter_outline(self, chapter_num: int) -> str:
        content = self._read_file("章节大纲索引.docx")
        pattern = rf"【第{chapter_num}章大纲】\s*\n(.*?)(?=\n【第\d+章大纲】|\Z)"
        match = re.search(pattern, content, re.DOTALL)
        return match.group(1).strip() if match else ""

    def read_progress(self) -> dict:
        content = self._read_file("写作进度.docx")
        progress = {}
        for line in content.strip().split("\n"):
            if "：" in line:
                key, value = line.split("：", 1)
                progress[key.strip("- ").strip()] = value.strip()
        return progress

    def update_progress(self, completed: int, next_chapter: int, status: str, context: str):
        content = self._format_progress(completed, next_chapter, status, context)
        self._write_file("写作进度.docx", content)

    def update_global_memory(self, updates: str):
        if not updates or updates.strip() == "无变化":
            return
        current = self.read_global_memory()
        self._write_file("全局记忆.docx", current + "\n\n" + updates)

    def update_clue_ledger(self, clue_changes: list):
        if not clue_changes:
            return
        current = self.read_clue_ledger()
        changes_text = "\n".join(f"- {c}" for c in clue_changes)
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        updated = current.replace(
            "【最后更新时间】",
            f"{changes_text}\n\n【最后更新时间】"
        )
        if "【最后更新时间】" not in updated:
            updated += f"\n\n【最后更新时间】{now}"
        self._write_file("线索台账.docx", updated)

    def update_recent_summaries(self, chapter_num: int, summary: str):
        current = self.read_recent_summaries()
        lines = [l for l in current.strip().split("\n") if l.strip() and not l.startswith("【")]
        lines.append(f"第{chapter_num}章：{summary}")
        if len(lines) > self.config.recent_summary_window:
            lines = lines[-self.config.recent_summary_window:]
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        content = "\n".join(lines) + f"\n\n【最后更新时间】{now}"
        self._write_file("近期摘要.docx", content)

    def _format_progress(self, completed: int, next_chapter: int, status: str, context: str = "") -> str:
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        return (
            f"- 已完成：{completed}章\n"
            f"- 下一章：第{next_chapter}章\n"
            f"- 状态：{status}\n"
            f"- 最后更新时间：{now}\n"
            f"- 当前任务上下文：{context}"
        )

    def _write_file(self, filename: str, content: str):
        filepath = os.path.join(self.mp, filename)
        doc = Document()
        doc.add_paragraph(content)
        doc.save(filepath)

    def _read_file(self, filename: str) -> str:
        filepath = os.path.join(self.mp, filename)
        if not os.path.exists(filepath):
            return ""
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)


class DocxHandler:

    def __init__(self, config: Config):
        self.config = config

    def save_chapter(self, chapter_num: int, content: str) -> str:
        filepath = os.path.join(self.config.chapter_path, f"第{chapter_num:03d}章.docx")
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)

        paragraphs = content.strip().split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(para_text)
            run.font.name = "宋体"
            run.font.size = Pt(12)
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.line_spacing = 1.5

        doc.save(filepath)
        return filepath

    def format_check(self, chapter_num: int, content: str) -> Tuple[bool, list]:
        issues = []
        filepath = os.path.join(self.config.chapter_path, f"第{chapter_num:03d}章.docx")

        if not os.path.exists(filepath):
            issues.append(f"文件未保存：{filepath}")

        char_count = len(content.replace("\n", "").replace(" ", ""))
        min_words = self.config.words_per_chapter - 300
        max_words = self.config.words_per_chapter + 300
        if char_count < min_words or char_count > max_words:
            issues.append(f"字数异常：{char_count}字，要求{min_words}-{max_words}字")

        paragraphs = [p.strip() for p in content.split("\n") if p.strip()]
        for i, p in enumerate(paragraphs):
            p_len = len(p.replace(" ", ""))
            if p_len < 50:
                issues.append(f"第{i+1}段过短：{p_len}字（要求50-250字）")
            elif p_len > 250:
                issues.append(f"第{i+1}段过长：{p_len}字（要求50-250字）")

        forbidden_patterns = [
            (r"---元数据---", "混入元数据"),
            (r"---元数据结束---", "混入元数据"),
            (r"摘要：", "混入摘要"),
            (r"线索变化：", "混入线索变化"),
        ]
        for pattern, desc in forbidden_patterns:
            if re.search(pattern, content):
                issues.append(f"正文纯度违规：{desc}")

        if "\n\n\n" in content:
            issues.append("存在连续空行")

        md_patterns = [
            (r"(?<!\w)\*{1,3}(?!\w)", "Markdown星号"),
            (r"^#{1,6}\s", "Markdown标题"),
            (r"^>\s", "Markdown引用"),
            (r"```", "Markdown代码块"),
        ]
        for pattern, desc in md_patterns:
            if re.search(pattern, content, re.MULTILINE):
                issues.append(f"格式违规：包含{desc}")

        passed = len(issues) == 0
        return passed, issues

    def fix_format(self, content: str) -> str:
        content = re.sub(r"\n{3,}", "\n\n", content)
        content = re.sub(r"(?<!\w)\*{1,3}(?!\w)", "", content)
        content = re.sub(r"^#{1,6}\s", "", content, flags=re.MULTILINE)
        content = re.sub(r"^>\s", "", content, flags=re.MULTILINE)
        content = content.replace("```", "")
        return content.strip()

    def save_review_record(self, filename: str, content: str):
        filepath = os.path.join(self.config.review_path, filename)
        doc = Document()
        doc.add_paragraph(content)
        doc.save(filepath)


class ReviewStandards:

    HARD_ERRORS = [
        (r"不是.*?而是", "禁止'不是...而是...'句式"),
        (r"——", "禁止破折号「——」"),
        (r"第\d+章|chapter\s+\d+", "禁止章节号指称"),
    ]

    REPORT_TERMS = [
        "核心动机", "信息落差", "认知共鸣", "锚定效应",
        "叙事张力", "情感锚点", "角色弧光", "戏剧冲突",
    ]

    TRANSITION_WORDS = [
        "仿佛", "忽然", "竟然", "猛地", "猛然", "不禁", "宛如",
    ]

    PREACHY_WORDS = [
        "显然", "毋庸置疑", "不言而喻", "众所周知",
    ]

    SHOCK_PHRASES = [
        "全场震惊", "众人惊呆", "所有人都愣住了", "一片哗然",
    ]

    @staticmethod
    def check_hard_errors(content: str) -> list:
        issues = []
        for pattern, desc in ReviewStandards.HARD_ERRORS:
            if re.search(pattern, content):
                issues.append(f"硬性错误：{desc}")

        for term in ReviewStandards.REPORT_TERMS:
            if term in content:
                issues.append(f"硬性错误：包含分析报告术语「{term}」")

        md_patterns = [
            (r"(?<!\w)\*{1,3}(?!\w)", "Markdown星号"),
            (r"^#{1,6}\s", "Markdown标题"),
            (r"^>\s", "Markdown引用"),
            (r"```", "Markdown代码块"),
        ]
        for pattern, desc in md_patterns:
            if re.search(pattern, content, re.MULTILINE):
                issues.append(f"硬性错误：包含{desc}")

        return issues

    @staticmethod
    def check_warnings(content: str) -> list:
        issues = []
        char_count = len(content.replace("\n", "").replace(" ", ""))

        transition_count = sum(content.count(w) for w in ReviewStandards.TRANSITION_WORDS)
        max_transitions = max(1, char_count // 2000)
        if transition_count > max_transitions:
            issues.append(f"警告：转折词密度过高（{transition_count}次/{char_count}字，允许{max_transitions}次）")

        for word in ReviewStandards.PREACHY_WORDS:
            if word in content:
                issues.append(f"警告：包含说教词「{word}」")

        for phrase in ReviewStandards.SHOCK_PHRASES:
            if phrase in content:
                issues.append(f"警告：包含集体震惊套话「{phrase}」")

        sentences = re.split(r"[。！？]", content)
        consecutive_le = 0
        for s in sentences:
            if "了" in s:
                consecutive_le += 1
                if consecutive_le >= 6:
                    issues.append("警告：连续6句以上含「了」字")
                    break
            else:
                consecutive_le = 0

        words = re.findall(r"[\u4e00-\u9fff]{2,4}", content)
        word_count = {}
        for w in words:
            word_count[w] = word_count.get(w, 0) + 1
        for w, c in word_count.items():
            if c > 3 and len(w) <= 2:
                issues.append(f"警告：高频词「{w}」出现{c}次")

        return issues


class PromptBuilder:

    def __init__(self, config: Config):
        self.config = config

    def build_material_confirmation_prompt(
        self,
        story_bible: str,
        chapter_outlines: str,
        total_chapters: int,
        words_per_chapter: int,
    ) -> Tuple[str, str]:
        system = (
            '你是一个小说项目审核AI。用户要启动一个长篇小说的自动写作系统，'
            '你需要在启动前检查用户提供的材料是否完整、是否有矛盾、是否有模糊不清的地方。\n\n'
            '【检查项】\n'
            '1. 故事圣经是否包含：世界观设定、核心角色档案（姓名/年龄/身份/性格/关键特征/初始状态）、'
            '主线剧情框架、支线剧情框架\n'
            '2. 章节大纲是否覆盖全部章节，每章是否有明确的剧情要点\n'
            '3. 故事圣经和章节大纲之间是否有逻辑矛盾\n'
            '4. 是否有模糊不清需要澄清的地方\n'
            '5. 章节数量是否与大纲匹配\n\n'
            '【输出格式】\n'
            '严格按以下JSON格式输出：\n'
            '{"passed": true/false, "issues": ["问题描述"], "suggestions": ["建议补充的内容"]}'
        )

        user = (
            f"【总章节数】{total_chapters}\n"
            f"【每章目标字数】{words_per_chapter}\n\n"
            f"【故事圣经】\n{story_bible}\n\n"
            f"【章节大纲】\n{chapter_outlines}"
        )

        return system, user

    def build_writing_prompt(
        self,
        global_memory: str,
        clue_ledger: str,
        recent_summaries: str,
        chapter_outline: str,
        chapter_num: int,
        feedback: str = None,
    ) -> Tuple[str, str]:
        style_requirement = ""
        if self.config.writing_style:
            style_requirement = f"\n【写作风格要求】\n{self.config.writing_style}\n"

        special_requirement = ""
        if self.config.special_requirements:
            special_requirement = f"\n【特殊要求】\n{self.config.special_requirements}\n"

        custom_review = ""
        if self.config.review_standards:
            custom_review = f"\n【自定义审核标准】\n{self.config.review_standards}\n"

        system = (
            '你是一个小说写作AI。你只需要完成一个任务：根据提供的材料写一章正文。\n\n'
            '【写作要求】\n'
            f'1. 创作约{self.config.words_per_chapter}字正文\n'
            '2. 不得与记忆系统中的任何信息矛盾\n'
            '3. 自然收尾，留白或场景转换，禁止总结性段落和AI式结尾\n'
            f'{style_requirement}'
            f'{special_requirement}'
            f'{custom_review}'
            '\n'
            '【违禁规则 - 硬性】\n'
            '- 禁止「不是...而是...」句式\n'
            '- 禁止破折号「——」\n'
            '- 禁止章节号指称（第X章、chapter X）\n'
            '- 禁止分析报告术语（核心动机、信息落差、认知共鸣、锚定效应等）\n'
            '- 禁止未设定的人物/地名\n'
            '- 禁止真实地名（除非大纲明确）\n'
            '- 禁止违背现实逻辑的情节\n'
            '- 禁止在正文中使用Markdown格式符号（*、#、-、_、```、> 等）\n\n'
            '【违禁规则 - 警告】\n'
            f'- 转折词密度：每{self.config.transition_word_interval}字不超过1次'
            '（仿佛、忽然、竟然、猛地、猛然、不禁、宛如）\n'
            '- 高疲劳词：同词每章只出现1次\n'
            '- 禁止元叙事/编剧旁白\n'
            '- 禁止说教词（显然、毋庸置疑、不言而喻、众所周知）\n'
            '- 禁止集体震惊套话（全场震惊、众人惊呆）\n'
            '- 禁止连续6句以上含「了」字\n'
            '- 段落控制在50-250字\n'
            '- 章节标题不得与已有章节重复或高度相似\n\n'
            '【输出格式 - 严格遵守】\n'
            '正文结束后，必须另起一段，严格按以下格式输出元数据。'
            '元数据必须包含以下三个部分，缺一不可：\n\n'
            '---元数据---\n'
            '摘要：（100字以内的本章内容摘要）\n'
            '线索变化：\n'
            '- （线索名称）：从（旧状态）变为（新状态）\n'
            '- （线索名称）：无变化\n'
            '建议更新：（如有角色状态变化等需要更新全局记忆的内容，没有则写「无」）\n'
            '---元数据结束---\n\n'
            '注意：\n'
            '1. 元数据仅用于系统更新记忆文件，不会写入正文\n'
            '2. 线索变化必须用「从xxx变为xxx」格式，每条一行\n'
            '3. 必须以---元数据---开头，以---元数据结束---结尾'
        )

        user = f"【世界观规则】\n{global_memory}\n\n"
        user += f"【线索台账】\n{clue_ledger}\n\n"
        user += f"【近期摘要】\n{recent_summaries}\n\n"
        user += f"【第{chapter_num}章大纲】\n{chapter_outline}\n\n"
        user += f"请创作第{chapter_num}章正文。"

        if feedback:
            user += f"\n\n【上一次审核反馈，请务必修正】\n{feedback}"

        return system, user

    def build_review_prompt(
        self,
        content: str,
        clue_ledger: str,
        recent_summaries: str,
        chapter_num: int,
    ) -> Tuple[str, str]:
        custom_review = ""
        if self.config.review_standards:
            custom_review = f"\n【自定义审核标准】\n{self.config.review_standards}\n"

        system = (
            '你是一个小说审核AI。请严格检查以下章节内容。\n\n'
            '【检查项】\n'
            '1. 与近期摘要中上一章结尾是否无缝衔接\n'
            '2. 本章内容是否与线索台账矛盾\n'
            '3. 人物状态是否与全局记忆一致\n'
            '4. 情节是否符合大纲要求\n'
            '5. 是否存在硬性违禁项\n'
            '6. 是否存在警告项\n'
            f'{custom_review}'
            '\n'
            '【severity严重程度说明】\n'
            '- hard：存在硬性错误，必须重写\n'
            '- warning：存在警告项，需要修改\n'
            '- none：无问题\n\n'
            '【输出格式】\n'
            '严格按以下JSON格式输出，不要输出其他内容：\n'
            '{"passed": true/false, "issues": ["问题描述"], "severity": "hard/warning/none"}'
        )

        user = f"【线索台账】\n{clue_ledger}\n\n"
        user += f"【近期摘要】\n{recent_summaries}\n\n"
        user += f"【待审核章节 - 第{chapter_num}章】\n{content}"

        return system, user

    def build_memory_init_prompt(self, story_bible: str) -> Tuple[str, str]:
        system = (
            '你是一个小说设定分析AI。请从故事圣经中提取以下信息，严格按JSON格式输出：\n\n'
            '{\n'
            '  "global_memory": "世界观核心规则（不超过500字）\\n\\n核心角色档案（每个角色：姓名、年龄、身份、性格、关键特征、初始状态）",\n'
            '  "clue_ledger": "角色状态追踪\\n物品道具追踪\\n伏笔追踪\\n时间线\\n\\n【最后更新时间】初始化",\n'
            '  "outline_index": "按章拆分的大纲，每章格式：\\n【第X章大纲】\\n内容"\n'
            '}'
        )
        user = f"【故事圣经】\n{story_bible}"
        return system, user

    def build_chapter_analysis_prompt(self, chapter_content: str, chapter_num: int) -> Tuple[str, str]:
        system = (
            '你是一个小说分析AI。请分析以下章节内容，提取关键信息，严格按JSON格式输出：\n'
            '{"summary": "100字以内摘要", '
            '"character_states": ["角色当前状态"], '
            '"clue_changes": ["线索变化"], '
            '"items": ["物品状态变化"], '
            '"foreshadowing": ["伏笔状态"]}'
        )
        user = f"【第{chapter_num}章内容】\n{chapter_content}"
        return system, user


class NovelForge:

    def __init__(self, config: Config):
        self.config = config
        self.ai = AIInterface(config)
        self.memory = MemorySystem(config)
        self.docx = DocxHandler(config)
        self.prompts = PromptBuilder(config)

    def start(self, mode: str, **kwargs):
        if mode == "new":
            self._start_new(**kwargs)
        elif mode == "resume":
            self._start_resume()
        elif mode == "import":
            self._start_import(**kwargs)

    def _confirm_materials(
        self,
        story_bible: str,
        chapter_outlines: str,
    ) -> Tuple[bool, list]:
        print("[NovelForge] 正在确认材料完整性...")

        system, user = self.prompts.build_material_confirmation_prompt(
            story_bible, chapter_outlines,
            self.config.total_chapters, self.config.words_per_chapter
        )

        try:
            response = self.ai.generate(system, user)
            result = self._parse_json(response)

            if result:
                passed = result.get("passed", False)
                issues = result.get("issues", [])
                suggestions = result.get("suggestions", [])

                if passed:
                    print("[NovelForge] 材料确认通过")
                    return True, []
                else:
                    print(f"[NovelForge] 材料确认未通过，发现{len(issues)}个问题")
                    for i, issue in enumerate(issues, 1):
                        print(f"  {i}. {issue}")
                    if suggestions:
                        print(f"  建议：{suggestions}")
                    return False, issues
            else:
                print("[NovelForge] 材料确认AI返回格式异常，跳过确认继续")
                return True, []

        except Exception as e:
            print(f"[NovelForge] 材料确认调用失败：{e}，跳过确认继续")
            return True, []

    def _start_new(self, story_bible: str, chapter_outlines: str):
        print(f"[NovelForge] 项目：{self.config.project_name}")
        print(f"[NovelForge] 总章数：{self.config.total_chapters}，每章{self.config.words_per_chapter}字")

        passed, issues = self._confirm_materials(story_bible, chapter_outlines)
        if not passed:
            print("\n[NovelForge] 材料有问题，请补充后重新启动：")
            for i, issue in enumerate(issues, 1):
                print(f"  {i}. {issue}")
            print("\n补充完毕后，重新调用 start('new') 即可。")
            return

        print(f"[NovelForge] 初始化项目：{self.config.project_name}")
        self.config.setup_paths()

        self._save_original("故事圣经.docx", story_bible)
        self._save_original("章节大纲.docx", chapter_outlines)

        print("[NovelForge] AI正在分析故事圣经，生成记忆系统...")
        system, user = self.prompts.build_memory_init_prompt(story_bible)
        ai_response = self.ai.generate(system, user)
        memory_data = self._parse_json(ai_response)

        self.memory.initialize(story_bible, chapter_outlines)
        if memory_data:
            if "global_memory" in memory_data:
                self.memory._write_file("全局记忆.docx", memory_data["global_memory"])
            if "clue_ledger" in memory_data:
                self.memory._write_file("线索台账.docx", memory_data["clue_ledger"])
            if "outline_index" in memory_data:
                self.memory._write_file("章节大纲索引.docx", memory_data["outline_index"])

        print("[NovelForge] 记忆系统初始化完成")
        print("[NovelForge] 开始自动写作循环")
        self._main_loop(start_chapter=1)

    def _start_resume(self):
        print("[NovelForge] 检测到断点续写")
        self.config.setup_paths()

        progress = self.memory.read_progress()
        next_chapter = int(progress.get("下一章", "第1章").replace("第", "").replace("章", ""))
        print(f"[NovelForge] 从第{next_chapter}章继续写作")

        self._main_loop(start_chapter=next_chapter)

    def _start_import(self, story_bible: str, chapter_outlines: str,
                      existing_chapters: dict, start_chapter: int):
        print(f"[NovelForge] 导入模式：从第{start_chapter}章开始续写")

        passed, issues = self._confirm_materials(story_bible, chapter_outlines)
        if not passed:
            print("\n[NovelForge] 材料有问题，请补充后重新启动：")
            for i, issue in enumerate(issues, 1):
                print(f"  {i}. {issue}")
            return

        self.config.setup_paths()

        self._save_original("故事圣经.docx", story_bible)
        self._save_original("章节大纲.docx", chapter_outlines)

        self.memory.initialize(story_bible, chapter_outlines)

        total_import = start_chapter - 1

        print(f"[NovelForge] 开始逐章导入，共{total_import}章")
        for i in range(1, start_chapter):
            chapter_content = existing_chapters.get(i, "")
            if not chapter_content:
                print(f"[NovelForge] 警告：第{i}章内容缺失，跳过")
                continue

            print(f"[NovelForge] 导入第{i}章...")
            system, user = self.prompts.build_chapter_analysis_prompt(chapter_content, i)
            analysis = self.ai.generate(system, user)
            analysis_data = self._parse_json(analysis)

            if analysis_data:
                if "summary" in analysis_data:
                    self.memory.update_recent_summaries(i, analysis_data["summary"])
                if "clue_changes" in analysis_data:
                    self.memory.update_clue_ledger(analysis_data["clue_changes"])

            self.docx.save_chapter(i, chapter_content)

        self.memory.update_progress(total_import, start_chapter, "写作中",
                                    f"导入完成，即将开始第{start_chapter}章写作")
        print(f"[NovelForge] 导入完成，开始写作第{start_chapter}章")

        self._main_loop(start_chapter=start_chapter)

    def _main_loop(self, start_chapter: int):
        for chapter_num in range(start_chapter, self.config.total_chapters + 1):
            print(f"\n{'='*60}")
            print(f"[NovelForge] 开始第{chapter_num}章")
            print(f"{'='*60}")

            success = self._write_and_save_chapter(chapter_num)

            if not success:
                print(f"[NovelForge] 第{chapter_num}章处理失败，流程暂停")
                self.memory.update_progress(
                    chapter_num - 1, chapter_num, "暂停-需要人工介入",
                    f"第{chapter_num}章异常处理全部失败"
                )
                return

            if chapter_num % self.config.batch_review_interval == 0:
                print(f"[NovelForge] 第{chapter_num}章完成，触发阶段审核")
                self._batch_review(chapter_num)

        print(f"\n[NovelForge] 全部{self.config.total_chapters}章写作完成，开始终审")
        self._final_review()

    def _write_and_save_chapter(self, chapter_num: int) -> bool:
        global_memory = self.memory.read_global_memory()
        clue_ledger = self.memory.read_clue_ledger()
        recent_summaries = self.memory.read_recent_summaries()
        chapter_outline = self.memory.read_chapter_outline(chapter_num)

        print(f"[写作] 调用AI生成第{chapter_num}章...")
        content, metadata = self._call_ai_write(
            chapter_num, global_memory, clue_ledger,
            recent_summaries, chapter_outline
        )

        if not content:
            return False

        print(f"[审核] 审核第{chapter_num}章...")
        passed, issues, severity = self._call_ai_review(
            content, clue_ledger, recent_summaries, chapter_num
        )

        retry_count = 0
        while not passed and retry_count < self.config.chapter_retry_max:
            retry_count += 1
            print(f"[审核] 未通过（severity={severity}，第{retry_count}次），重写...")
            feedback = "\n".join(issues)
            content, metadata = self._call_ai_write(
                chapter_num, global_memory, clue_ledger,
                recent_summaries, chapter_outline, feedback=feedback
            )
            if content:
                passed, issues, severity = self._call_ai_review(
                    content, clue_ledger, recent_summaries, chapter_num
                )

        if not passed:
            print(f"[异常] 第{chapter_num}章审核3次未通过（severity={severity}），进入异常处理")
            content, metadata = self._handle_failure(
                chapter_num, global_memory, clue_ledger,
                recent_summaries, chapter_outline
            )
            if not content:
                return False

        print(f"[保存] 保存第{chapter_num}章.docx")
        self.docx.save_chapter(chapter_num, content)

        format_ok, format_issues = self.docx.format_check(chapter_num, content)
        if not format_ok:
            print(f"[格式] 格式检查未通过，自动修复：{format_issues}")
            content = self.docx.fix_format(content)
            self.docx.save_chapter(chapter_num, content)

        if metadata:
            self.memory.update_recent_summaries(chapter_num, metadata.get("summary", ""))
            self.memory.update_clue_ledger(metadata.get("clue_changes", []))
            if metadata.get("suggested_updates"):
                self.memory.update_global_memory("\n".join(metadata["suggested_updates"]))

        self.memory.update_progress(
            chapter_num, chapter_num + 1, "写作中",
            f"第{chapter_num}章已完成，即将开始第{chapter_num + 1}章"
        )

        del content

        print(f"[完成] 第{chapter_num}章完成")
        return True

    def _call_ai_write(
        self,
        chapter_num: int,
        global_memory: str,
        clue_ledger: str,
        recent_summaries: str,
        chapter_outline: str,
        feedback: str = None,
    ) -> Tuple[Optional[str], Optional[dict]]:
        try:
            system, user = self.prompts.build_writing_prompt(
                global_memory, clue_ledger, recent_summaries,
                chapter_outline, chapter_num, feedback
            )
            response = self.ai.generate(system, user)
            content, metadata = self._split_content_and_metadata(response)
            return content, metadata
        except Exception as e:
            print(f"[错误] AI写作调用失败：{e}")
            return None, None

    def _call_ai_review(
        self,
        content: str,
        clue_ledger: str,
        recent_summaries: str,
        chapter_num: int,
    ) -> Tuple[bool, list, str]:
        hard_issues = ReviewStandards.check_hard_errors(content)
        if hard_issues:
            return False, hard_issues, "hard"

        warning_issues = ReviewStandards.check_warnings(content)

        try:
            system, user = self.prompts.build_review_prompt(
                content, clue_ledger, recent_summaries, chapter_num
            )
            response = self.ai.generate(system, user)
            review_data = self._parse_json(response)

            if review_data:
                passed = review_data.get("passed", False)
                ai_issues = review_data.get("issues", [])
                severity = review_data.get("severity", "none")

                if severity == "hard":
                    passed = False

                all_issues = warning_issues + ai_issues
                return passed, all_issues, severity
            else:
                local_severity = "warning" if warning_issues else "none"
                return len(hard_issues) == 0, warning_issues, local_severity

        except Exception as e:
            print(f"[错误] AI审核调用失败：{e}")
            local_severity = "warning" if warning_issues else "none"
            return len(hard_issues) == 0, warning_issues, local_severity

    def _handle_failure(
        self,
        chapter_num: int,
        global_memory: str,
        clue_ledger: str,
        recent_summaries: str,
        chapter_outline: str,
    ) -> Tuple[Optional[str], Optional[dict]]:
        print("[异常处理] 第一轮：原始重试")
        for i in range(3):
            print(f"[异常处理] 第一轮重试 {i+1}/3")
            content, metadata = self._call_ai_write(
                chapter_num, global_memory, clue_ledger,
                recent_summaries, chapter_outline
            )
            if content:
                passed, _, _ = self._call_ai_review(
                    content, clue_ledger, recent_summaries, chapter_num
                )
                if passed:
                    return content, metadata

        print("[异常处理] 第二轮：精简内容重试")
        minimal_outline = chapter_outline[:500]
        minimal_clues = clue_ledger[:300]
        minimal_summaries = "\n".join(
            recent_summaries.strip().split("\n")[-2:]
        )
        content, metadata = self._call_ai_write(
            chapter_num, global_memory[:300],
            minimal_clues, minimal_summaries, minimal_outline
        )
        if content:
            return content, metadata

        print("[异常处理] 第三轮：等待10分钟后重试")
        time.sleep(600)
        content, metadata = self._call_ai_write(
            chapter_num, global_memory, clue_ledger,
            recent_summaries, chapter_outline
        )
        if content:
            passed, _, _ = self._call_ai_review(
                content, clue_ledger, recent_summaries, chapter_num
            )
            if passed:
                return content, metadata

        print("[异常处理] 第四轮：全部失败，暂停")
        failure_record = (
            f"第{chapter_num}章异常排查记录\n"
            f"时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"第一轮：原始重试3次，全部失败\n"
            f"第二轮：精简内容重试，失败\n"
            f"第三轮：等待10分钟冷却后重试，失败\n"
            f"第四轮：暂停，等待人工介入"
        )
        self.docx.save_review_record(f"第{chapter_num:03d}章_异常排查.docx", failure_record)
        return None, None

    def _rewrite_chapter(self, chapter_num: int) -> bool:
        global_memory = self.memory.read_global_memory()
        clue_ledger = self.memory.read_clue_ledger()
        recent_summaries = self.memory.read_recent_summaries()
        chapter_outline = self.memory.read_chapter_outline(chapter_num)

        content, metadata = self._call_ai_write(
            chapter_num, global_memory, clue_ledger,
            recent_summaries, chapter_outline
        )

        if not content:
            print(f"[重写] 第{chapter_num}章重写失败")
            return False

        self.docx.save_chapter(chapter_num, content)

        format_ok, format_issues = self.docx.format_check(chapter_num, content)
        if not format_ok:
            content = self.docx.fix_format(content)
            self.docx.save_chapter(chapter_num, content)

        if metadata:
            self.memory.update_recent_summaries(chapter_num, metadata.get("summary", ""))
            self.memory.update_clue_ledger(metadata.get("clue_changes", []))
            if metadata.get("suggested_updates"):
                self.memory.update_global_memory("\n".join(metadata["suggested_updates"]))

        del content
        print(f"[重写] 第{chapter_num}章重写完成")
        return True

    def _batch_review(self, up_to_chapter: int):
        print(f"\n[阶段审核] 对第1-{up_to_chapter}章进行台账审查")

        clue_ledger = self.memory.read_clue_ledger()
        global_memory = self.memory.read_global_memory()
        recent_summaries = self.memory.read_recent_summaries()

        system = (
            '你是一个小说审核编辑。请严格审查以下线索台账，检查：\n'
            '1. 每条线索的演变是否自洽\n'
            '2. 人物状态是否合理\n'
            '3. 伏笔是否有遗忘\n'
            '4. 时间线是否混乱\n\n'
            '只有发现可疑项时，才需要回溯相关章节原文核实。\n\n'
            '按JSON格式输出：\n'
            '{"passed": true/false, "issues": ["问题描述"], '
            '"fix_suggestions": ["修复建议"], '
            '"chapters_to_rewrite": [需要重写的章节号列表，没有则为空数组]}'
        )
        user = (
            f"【全局记忆】\n{global_memory}\n\n"
            f"【线索台账】\n{clue_ledger}\n\n"
            f"【近期摘要】\n{recent_summaries}"
        )

        try:
            response = self.ai.generate(system, user)
            review_data = self._parse_json(response)

            if review_data and not review_data.get("passed", True):
                issues = review_data.get("issues", [])
                fixes = review_data.get("fix_suggestions", [])
                chapters_to_rewrite = review_data.get("chapters_to_rewrite", [])

                review_report = (
                    f"第{up_to_chapter}章阶段审核报告\n"
                    f"时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                    f"发现问题：\n" + "\n".join(f"- {i}" for i in issues) +
                    f"\n修复建议：\n" + "\n".join(f"- {f}" for f in fixes) +
                    f"\n需要重写的章节：{chapters_to_rewrite}"
                )
                self.docx.save_review_record(
                    f"第{up_to_chapter}章阶段审核.docx", review_report
                )

                print(f"[阶段审核] 发现{len(issues)}个问题，需要重写{len(chapters_to_rewrite)}章")

                for chapter_num in chapters_to_rewrite:
                    print(f"[阶段审核] 重写第{chapter_num}章...")
                    self._rewrite_chapter(chapter_num)

                print("[阶段审核] 所有问题已修复，继续写作")
            else:
                print("[阶段审核] 通过")

        except Exception as e:
            print(f"[阶段审核] 审核调用失败：{e}")

    def _final_review(self):
        print("[终审] 开始最终审查")
        self._final_review_loop(remaining_rounds=self.config.final_review_max_rounds)

    def _final_review_loop(self, remaining_rounds: int):
        if remaining_rounds <= 0:
            print("[终审] 已重审3轮仍有问题，暂停等待人工介入")
            self.memory.update_progress(
                self.config.total_chapters, self.config.total_chapters,
                "暂停-终审未通过", "终审3轮仍有问题，等待人工介入"
            )
            return

        clue_ledger = self.memory.read_clue_ledger()
        global_memory = self.memory.read_global_memory()
        recent_summaries = self.memory.read_recent_summaries()

        system = (
            '你是一个小说终审编辑。请对整部小说的线索台账做最终通查。\n'
            '检查所有线索是否完整收束、人物弧线是否完成、是否有遗留问题。\n\n'
            '只有发现可疑项时，才需要回溯相关章节原文核实。\n\n'
            '按JSON格式输出：\n'
            '{"passed": true/false, "issues": ["问题描述"], '
            '"overall_quality": "评价", '
            '"chapters_to_rewrite": [需要重写的章节号列表，没有则为空数组]}'
        )
        user = (
            f"【全局记忆】\n{global_memory}\n\n"
            f"【线索台账】\n{clue_ledger}\n\n"
            f"【近期摘要】\n{recent_summaries}"
        )

        try:
            response = self.ai.generate(system, user)
            review_data = self._parse_json(response)

            if review_data:
                passed = review_data.get("passed", False)
                issues = review_data.get("issues", [])
                chapters_to_rewrite = review_data.get("chapters_to_rewrite", [])

                if not passed:
                    round_num = self.config.final_review_max_rounds - remaining_rounds + 1
                    report = (
                        f"终审报告（第{round_num}轮）\n"
                        f"时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                        f"总章数：{self.config.total_chapters}\n"
                        f"通过：False\n"
                        f"问题：\n" + "\n".join(f"- {i}" for i in issues) +
                        f"\n整体评价：{review_data.get('overall_quality', '')}"
                    )
                    self.docx.save_review_record(
                        f"终审报告_第{round_num}轮.docx", report
                    )

                    print(f"[终审] 第{round_num}轮：发现{len(issues)}个问题，需要重写{len(chapters_to_rewrite)}章")

                    for chapter_num in chapters_to_rewrite:
                        print(f"[终审] 重写第{chapter_num}章...")
                        self._rewrite_chapter(chapter_num)

                    print("[终审] 问题已修复，重新终审...")
                    self._final_review_loop(remaining_rounds - 1)
                    return
                else:
                    report = (
                        f"终审报告\n"
                        f"时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                        f"总章数：{self.config.total_chapters}\n"
                        f"通过：True\n"
                        f"问题：无\n"
                        f"整体评价：{review_data.get('overall_quality', '')}"
                    )
                    self.docx.save_review_record("终审报告.docx", report)
                    print("[终审] 通过")

        except Exception as e:
            print(f"[终审] 终审调用失败：{e}")
            return

        self.memory.update_progress(
            self.config.total_chapters, self.config.total_chapters,
            "已完成", f"全部{self.config.total_chapters}章写作完成，终审完成"
        )
        print(f"\n[NovelForge] {self.config.project_name} 全部完成！")

    def _split_content_and_metadata(self, response: str) -> Tuple[str, dict]:
        metadata = {}
        content = response

        meta_start = response.find("---元数据---")
        meta_end = response.find("---元数据结束---")

        if meta_start != -1 and meta_end != -1:
            content = response[:meta_start].strip()
            meta_text = response[meta_start + len("---元数据---"):meta_end].strip()

            metadata = {"summary": "", "clue_changes": [], "suggested_updates": []}
            current_section = None

            for line in meta_text.split("\n"):
                line = line.strip()
                if not line:
                    continue

                if line.startswith("摘要："):
                    metadata["summary"] = line.replace("摘要：", "").strip()
                    current_section = "summary"
                elif line.startswith("线索变化："):
                    current_section = "clues"
                elif line.startswith("建议更新："):
                    value = line.replace("建议更新：", "").strip()
                    if value and value != "无":
                        metadata["suggested_updates"].append(value)
                    current_section = "updates"
                elif line.startswith("- ") and current_section == "clues":
                    metadata["clue_changes"].append(line[2:].strip())

        return content, metadata

    def _parse_json(self, text: str) -> Optional[dict]:
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass

        match = re.search(r"```(?:json)?\s*\n?(.*?)\n?```", text, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(1))
            except json.JSONDecodeError:
                pass

        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1:
            try:
                return json.loads(text[start:end + 1])
            except json.JSONDecodeError:
                pass

        return None

    def _save_original(self, filename: str, content: str):
        filepath = os.path.join(self.config.original_path, filename)
        doc = Document()
        doc.add_paragraph(content)
        doc.save(filepath)


def main():
    print("=" * 60)
    print("  NovelForge v2.0.6 - 长篇小说全自动写作系统")
    print("=" * 60)

    config = Config()

    api_loaded = config.auto_load_api_config()

    if not api_loaded:
        print("\n请手动配置AI接口：")
        config.ai_base_url = input("API Base URL: ").strip()
        config.ai_api_key = input("API Key: ").strip()
        config.ai_model = input("Model Name: ").strip()

    print("\n请配置项目：")
    config.project_name = input("项目名称（小说名）: ").strip()
    config.total_chapters = int(input("总章节数: ").strip())
    words_input = input("每章目标字数（默认2500）: ").strip()
    config.words_per_chapter = int(words_input) if words_input else 2500

    print("\n可选材料（直接回车跳过）：")
    style_input = input("写作风格要求（文风、人称、叙事视角等）: ").strip()
    config.writing_style = style_input if style_input else ""
    special_input = input("其他特殊要求: ").strip()
    config.special_requirements = special_input if special_input else ""
    review_input = input("自定义审核标准（为空则使用内置标准）: ").strip()
    config.review_standards = review_input if review_input else ""

    forge = NovelForge(config)

    print("\n请选择模式：")
    print("1. 开始写（全新写作）")
    print("2. 继续写（断点续写）")
    print("3. 从第X章开始续写（导入模式）")

    choice = input("\n输入选项编号: ").strip()

    if choice == "1":
        print("\n请提供故事圣经（输入END结束）：")
        story_bible = _read_multiline()
        print("\n请提供章节大纲（输入END结束）：")
        chapter_outlines = _read_multiline()
        forge.start("new", story_bible=story_bible, chapter_outlines=chapter_outlines)

    elif choice == "2":
        forge.start("resume")

    elif choice == "3":
        start_chapter = int(input("从第几章开始续写: ").strip())
        print("\n请提供故事圣经（输入END结束）：")
        story_bible = _read_multiline()
        print("\n请提供章节大纲（输入END结束）：")
        chapter_outlines = _read_multiline()

        existing_chapters = {}
        print("\n请提供已完成的章节内容")
        print("（格式：先输入章节数字，再输入内容，输入0结束）")
        while True:
            num = input("\n章节号（输入0结束）: ").strip()
            if num == "0":
                break
            print(f"请输入第{num}章内容（输入END结束）：")
            content = _read_multiline()
            existing_chapters[int(num)] = content

        forge.start("import", story_bible=story_bible,
                    chapter_outlines=chapter_outlines,
                    existing_chapters=existing_chapters,
                    start_chapter=start_chapter)


def _read_multiline() -> str:
    lines = []
    while True:
        line = input()
        if line.strip() == "END":
            break
        lines.append(line)
    return "\n".join(lines)


def run_from_code():
    config = Config()
    config.auto_load_api_config()
    config.project_name = "我的小说"
    config.total_chapters = 500
    config.words_per_chapter = 2500
    config.writing_style = "第三人称，暗黑风格，节奏紧凑"

    forge = NovelForge(config)

    story_bible = Path("story_bible.txt").read_text(encoding="utf-8")
    chapter_outlines = Path("chapter_outlines.txt").read_text(encoding="utf-8")

    forge.start("new", story_bible=story_bible, chapter_outlines=chapter_outlines)


if __name__ == "__main__":
    main()
