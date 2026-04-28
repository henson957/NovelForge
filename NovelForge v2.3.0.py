#!/usr/bin/env python3
"""
NovelForge v2.3.0 - 长篇小说全自动写作系统

融合 v2.2.5 流程框架 + v2.2.2 质量保障体系 + InkOS 写作方法论。

变更日志：
  v2.3.0 - 融合 v2.2.2 的 InkOS 写作方法论、上下文精简、AI 痕迹检测、
           18 维度审计、信息边界、资源追踪、情感弧线、崩溃保护
         - 新增 ContextBuilder：按相关性精简上下文，避免 token 爆炸
         - 新增 AITellDetector：5 维度 AI 痕迹统计检测
         - 新增 PostWriteValidator：硬规则自动 spot-fix
         - 新增 AuditSystem：18 维度 3 层审计
         - 新增 character_matrix.md：角色信息边界追踪
         - 新增 resource_ledger.md：资源/道具数量追踪
         - 新增 emotional_arcs.md：角色情感弧线追踪
         - 新增 atomic_write：崩溃保护
         - 新增 InkOS 写作方法论注入 PromptBuilder
         - 新增审计→修订循环
         - 保留 v2.2.5 的材料确认、4 轮异常处理、阶段审核、终审、截断检测、导入模式
  v2.2.5 - 基础流程框架（材料确认、写作循环、审核、异常处理、阶段审核、终审）
  v2.2.2 - InkOS 方法论、AI 检测、审计系统、真相文件

依赖：
  pip install openai python-docx
"""

import os
import re
import json
import math
import time
from datetime import datetime
from typing import Tuple, Optional, Dict, List
from collections import Counter

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openai import OpenAI


# ============================================================
# 工具函数
# ============================================================

def atomic_write(filepath: str, content: str, as_docx: bool = False):
    """原子性写入：先写 .tmp，成功后替换，防止崩溃导致文件损坏"""
    tmp_path = filepath + ".tmp"
    if as_docx:
        doc = Document()
        doc.add_paragraph(content)
        doc.save(tmp_path)
    else:
        with open(tmp_path, "w", encoding="utf-8") as f:
            f.write(content)
    os.replace(tmp_path, filepath)


# ============================================================
# AI 痕迹检测（5 维度，纯统计）
# ============================================================

HEDGE_WORDS = ["似乎", "可能", "或许", "大概", "某种程度上", "一定程度上", "在某种意义上"]
TRANSITION_WORDS = ["然而", "不过", "与此同时", "另一方面", "尽管如此", "话虽如此", "但值得注意的是"]
FATIGUE_WORDS = ["仿佛", "忽然", "竟然", "猛地", "猛然", "不禁", "宛如", "随即", "旋即", "霎时", "顿时"]


class AITellDetector:

    @staticmethod
    def analyze(content: str) -> List[Dict]:
        issues = []
        total_chars = len(content)

        # 维度 1：段落等长
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', content) if p.strip()]
        if len(paragraphs) >= 3:
            lengths = [len(p) for p in paragraphs]
            mean = sum(lengths) / len(lengths)
            if mean > 0:
                cv = math.sqrt(sum((l - mean) ** 2 for l in lengths) / len(lengths)) / mean
                if cv < 0.15:
                    issues.append({"severity": "warning", "category": "段落等长",
                        "description": f"变异系数{cv:.3f}（阈值<0.15）",
                        "suggestion": "增加段落长度差异"})

        # 维度 2：套话密度
        if total_chars > 0:
            density = sum(content.count(w) for w in HEDGE_WORDS) / (total_chars / 1000)
            if density > 3:
                issues.append({"severity": "warning", "category": "套话密度",
                    "description": f"{density:.1f}次/千字（阈值>3）",
                    "suggestion": "用确定性叙述替代模糊表达"})

        # 维度 3：公式化转折
        trans = {w: content.count(w) for w in TRANSITION_WORDS if content.count(w) >= 3}
        if trans:
            issues.append({"severity": "warning", "category": "公式化转折",
                "description": "、".join(f'"{w}"×{c}' for w, c in trans.items()),
                "suggestion": "换用动作切入、时间跳跃、视角切换"})

        # 维度 4：列表式结构
        sentences = [s.strip() for s in re.split(r'[。！？\n]', content) if len(s.strip()) > 2]
        if len(sentences) >= 3:
            mc, c = 1, 1
            for i in range(1, len(sentences)):
                if sentences[i - 1][:2] == sentences[i][:2]:
                    c += 1; mc = max(mc, c)
                else:
                    c = 1
            if mc >= 3:
                issues.append({"severity": "info", "category": "列表式结构",
                    "description": f"{mc}句连续相同开头", "suggestion": "变换句式开头"})

        # 维度 5：高疲劳词
        if total_chars > 0:
            threshold = max(1, total_chars // 3000)
            hits = {w: content.count(w) for w in FATIGUE_WORDS if content.count(w) > threshold}
            if hits:
                issues.append({"severity": "warning", "category": "高疲劳词",
                    "description": "、".join(f'"{w}"×{c}' for w, c in hits.items()),
                    "suggestion": "改用具体动作或感官描写"})

        return issues


# ============================================================
# PostWriteValidator - 写后验证器
# ============================================================

class PostWriteValidator:

    FORBIDDEN_TERMS = [
        "核心动机", "信息边界", "信息落差", "核心风险", "利益最大化",
        "当前处境", "认知共鸣", "锚定效应", "沉没成本", "情绪缺口",
    ]

    @staticmethod
    def spot_fix(content: str) -> Tuple[str, List[str]]:
        fixes = []

        if "——" in content:
            content = content.replace("——", "，")
            fixes.append("破折号已替换为逗号")

        for term in PostWriteValidator.FORBIDDEN_TERMS:
            if term in content:
                content = content.replace(term, "...")
                fixes.append(f"术语'{term}'已标记")

        nb = re.findall(r'不是[^，。！？\n]{1,20}[，,]而?是[^，。！？\n]{1,20}', content)
        if nb:
            for m in nb:
                content = content.replace(m, f"[待改写:{m}]")
            fixes.append(f"'不是而是'{len(nb)}处已标记")

        ch_refs = set(re.findall(r'第\d+章|chapter\s*\d+', content, re.IGNORECASE))
        if ch_refs:
            for ref in ch_refs:
                content = content.replace(ref, "...")
            fixes.append(f"章节号指称{len(ch_refs)}处已清除")

        return content, fixes

    @staticmethod
    def check_repetition(chapter_num: int, content: str, summaries: str) -> List[Dict]:
        issues = []
        stop = {"一个", "这个", "那个", "他们", "她们", "我们", "什么", "怎么", "可以",
                "没有", "但是", "然后", "因为", "所以", "如果"}
        current_words = set(re.findall(r'[\u4e00-\u9fff]{2,4}', content)) - stop

        for ch in range(max(1, chapter_num - 5), chapter_num):
            m = re.search(rf'第{ch}章[^\n]*\n(.*?)(?=\n第\d+章|\n## |\Z)', summaries, re.DOTALL)
            if m:
                summary_words = set(re.findall(r'[\u4e00-\u9fff]{2,4}', m.group(1))) - stop
                if current_words and summary_words:
                    overlap = len(current_words & summary_words) / max(len(current_words), len(summary_words))
                    if overlap > 0.5:
                        issues.append({"severity": "warning", "category": "跨章词汇重复",
                            "description": f"与第{ch}章关键词重叠率{overlap:.1%}",
                            "suggestion": "检查是否有重复描写或情节"})

        sentences = [s.strip() for s in re.split(r'[。！？]', content) if len(s.strip()) > 10]
        seen: Dict[str, int] = {}
        for s in sentences:
            key = s[:15]
            seen[key] = seen.get(key, 0) + 1
        dups = [(k, c) for k, c in seen.items() if c >= 2]
        if dups:
            issues.append({"severity": "warning", "category": "章内句子重复",
                "description": f"{len(dups)}组相似句子", "suggestion": "删除或改写"})

        return issues


# ============================================================
# 18 维度审计系统
# ============================================================

class AuditSystem:

    FIXABLE = {"资源连续性", "信息越界", "违禁句式", "违禁符号", "违禁术语",
               "跨章词汇重复", "章内句子重复", "高疲劳词", "段落等长", "套话密度", "公式化转折"}

    def __init__(self, memory: 'MemorySystem', config: Config):
        self.memory = memory
        self.config = config

    def run_audit(self, chapter_num: int, content: str) -> List[Dict]:
        issues = []
        issues.extend(self._layer1(chapter_num, content))
        if chapter_num % 5 == 0:
            issues.extend(self._layer2(chapter_num))
        if chapter_num % 20 == 0:
            issues.extend(self._layer3(chapter_num))
        return issues

    def get_fixable_warnings(self, issues: List[Dict]) -> List[Dict]:
        return [i for i in issues if i["severity"] == "warning" and i["category"] in self.FIXABLE]

    # ---- Layer 1：每章（7 维度）----

    def _layer1(self, ch: int, content: str) -> List[Dict]:
        issues = []
        issues.extend(self._check_character_names(content))
        issues.extend(self._check_timeline(content, ch))
        issues.extend(self._check_resources(content))
        issues.extend(self._check_info_boundary(content))
        issues.extend(self._check_forbidden_patterns(content))
        issues.extend(PostWriteValidator.check_repetition(ch, content, self.memory.read_recent_summaries()))
        issues.extend(AITellDetector.analyze(content))
        return issues

    # ---- Layer 2：每 5 章（4 维度）----

    def _layer2(self, ch: int) -> List[Dict]:
        issues = []
        issues.extend(self._check_foreshadowing(ch))
        issues.extend(self._check_subplot_stagnation(ch))
        issues.extend(self._check_emotional_consistency())
        issues.extend(self._check_pacing(ch))
        return issues

    # ---- Layer 3：每 20 章（3 维度）----

    def _layer3(self, ch: int) -> List[Dict]:
        issues = []
        issues.extend(self._check_vocabulary_diversity(ch))
        issues.extend(self._check_character_balance(ch))
        issues.extend(self._check_numerical_consistency())
        return issues

    # ---- 维度实现 ----

    def _check_character_names(self, content: str) -> List[Dict]:
        tracking = self.memory.read_clue_ledger()
        if not tracking:
            return []
        known = set()
        for line in tracking.split("\n"):
            m = re.match(r'-\s*(\S+)', line.strip())
            if m:
                known.add(m.group(1))
        if not known:
            return []
        dialog_names = re.findall(r'["「](\S{2,4}?)[」"]\s*[,，。]?\s*(?:说|道|问|答|喊|笑|叹)', content)
        action_names = re.findall(r'(?:^|[。！？\n，])\s*(\S{2,4}?)(?:说|道|问|答|走|坐|站|怒|惊)', content)
        stop = {"他", "她", "我", "你", "我们", "他们", "对方", "那人", "这人", "一个"}
        unknown = set(dialog_names + action_names) - known - stop
        if unknown:
            return [{"severity": "info", "category": "角色名一致性",
                "description": f"未在台账中的角色名：{'、'.join(list(unknown)[:5])}",
                "suggestion": "确认是否为新角色"}]
        return []

    def _check_timeline(self, content: str, ch: int) -> List[Dict]:
        issues = []
        for num, unit in re.findall(r'(\d+)\s*(天|日|月|年|小时|分钟)\s*(?:后|前|以后|以前)', content):
            n = int(num)
            if unit == "年" and n > 10:
                issues.append({"severity": "warning", "category": "时间线连续性",
                    "description": f"第{ch}章出现{n}{unit}的时间跳跃", "suggestion": "需要铺垫"})
        return issues

    def _check_resources(self, content: str) -> List[Dict]:
        ledger = self.memory.read_resource_ledger()
        if not ledger:
            return []
        issues = []
        for m in re.finditer(r'(\S{2,8}?)\s*[:：]\s*(\d+)', ledger):
            name, expected = m.group(1), int(m.group(2))
            if name in content:
                actual_m = re.search(re.escape(name) + r'\s*[:：有剩]\s*(\d+)', content)
                if actual_m and int(actual_m.group(1)) != expected:
                    issues.append({"severity": "warning", "category": "资源连续性",
                        "description": f"'{name}'账本{expected}，正文{int(actual_m.group(1))}",
                        "suggestion": "核实资源数量"})
        return issues

    def _check_info_boundary(self, content: str) -> List[Dict]:
        matrix = self.memory.read_character_matrix()
        if not matrix:
            return []
        issues = []
        for line in matrix.split("\n"):
            if any(kw in line for kw in ("不知道", "未知", "未被告知")):
                parts = line.split("|", 2)
                if len(parts) >= 3:
                    info = parts[2].strip()
                    if len(info) >= 4 and info in content:
                        issues.append({"severity": "warning", "category": "信息越界",
                            "description": f"正文出现标记为'未知'的信息：'{info[:20]}'",
                            "suggestion": "检查信息越界"})
        return issues

    def _check_forbidden_patterns(self, content: str) -> List[Dict]:
        issues = []
        nb = re.findall(r'不是[^，。！？\n]{1,20}[，,]而?是[^，。！？\n]{1,20}', content)
        if nb:
            issues.append({"severity": "warning", "category": "违禁句式",
                "description": f"'不是而是'{len(nb)}处", "suggestion": "改用直述句"})
        if "——" in content:
            issues.append({"severity": "warning", "category": "违禁符号",
                "description": f"破折号{content.count('——')}处", "suggestion": "用逗号或句号"})
        terms = ["核心动机", "信息边界", "信息落差", "核心风险", "利益最大化", "当前处境"]
        found = [t for t in terms if t in content]
        if found:
            issues.append({"severity": "warning", "category": "违禁术语",
                "description": f"分析术语：{'、'.join(found)}", "suggestion": "替换为口语化表达"})
        return issues

    def _check_foreshadowing(self, ch: int) -> List[Dict]:
        text = self.memory.read_clue_ledger()
        if not text:
            return []
        issues = []
        for line in text.split("\n"):
            line = line.strip()
            if not line or "已回收" in line:
                continue
            ch_matches = re.findall(r'第(\d+)章', line)
            if ch_matches:
                gap = ch - max(int(c) for c in ch_matches)
                if gap >= 10:
                    issues.append({"severity": "warning", "category": "伏笔回收",
                        "description": f"'{line[:30]}'已{gap}章未推进", "suggestion": "回收或推进"})
        return issues

    def _check_subplot_stagnation(self, ch: int) -> List[Dict]:
        text = self.memory.read_clue_ledger()
        if not text:
            return []
        issues = []
        for line in text.split("\n"):
            line = line.strip()
            if not line or "已解决" in line or "已完成" in line:
                continue
            ch_matches = re.findall(r'第(\d+)章', line)
            if ch_matches:
                gap = ch - max(int(c) for c in ch_matches)
                if gap >= 5:
                    issues.append({"severity": "warning", "category": "支线停滞",
                        "description": f"'{line[:30]}'已{gap}章未推进", "suggestion": "推进或给出理由"})
        return issues

    def _check_emotional_consistency(self) -> List[Dict]:
        arcs = self.memory.read_emotional_arcs()
        if not arcs:
            return []
        positive = {"开心", "兴奋", "自信", "满足", "平静", "温暖", "希望", "释然"}
        negative = {"愤怒", "悲伤", "恐惧", "焦虑", "绝望", "痛苦", "压抑", "不安"}
        character_emotions: Dict[str, List[str]] = {}
        for line in arcs.split("\n"):
            line = line.strip()
            if not line or not line.startswith("-"):
                continue
            parts = re.split(r'[\|｜]', line.lstrip("- "))
            if len(parts) >= 2:
                name = parts[0].strip()
                emotion_field = parts[1].strip()
                emotion_words = [w for w in positive | negative if w in emotion_field]
                if emotion_words:
                    character_emotions.setdefault(name, []).append(emotion_words[0])
        issues = []
        for name, emotions in character_emotions.items():
            if len(emotions) >= 3:
                recent = emotions[-3:]
                flips = sum(1 for i in range(1, len(recent))
                           if (recent[i] in positive) != (recent[i - 1] in positive))
                if flips >= 2:
                    issues.append({"severity": "info", "category": "情感弧线",
                        "description": f"'{name}'情绪反复：{'→'.join(recent)}",
                        "suggestion": "情绪变化需要事件驱动"})
        return issues

    def _check_pacing(self, ch: int) -> List[Dict]:
        summaries = self.memory.read_recent_summaries()
        if not summaries:
            return []
        high = {"战斗", "冲突", "高潮", "追杀", "对决", "爆发", "危机", "对峙", "翻脸"}
        low = {"过渡", "日常", "修炼", "铺垫", "准备", "休息", "闲聊", "整理"}
        recent_types = []
        for line in summaries.split("\n"):
            m = re.search(r'第(\d+)章', line)
            if m and ch - 5 <= int(m.group(1)) <= ch:
                if any(w in line for w in high):
                    recent_types.append("high")
                elif any(w in line for w in low):
                    recent_types.append("low")
        issues = []
        if len(recent_types) >= 4:
            tail = recent_types[-4:]
            if all(t == "high" for t in tail):
                issues.append({"severity": "info", "category": "节奏张弛",
                    "description": f"最近{len(tail)}章全是高张力", "suggestion": "插入过渡章节"})
            elif all(t == "low" for t in tail):
                issues.append({"severity": "warning", "category": "节奏张弛",
                    "description": f"最近{len(tail)}章全是低张力", "suggestion": "推进冲突或释放爽点"})
        return issues

    def _check_vocabulary_diversity(self, ch: int) -> List[Dict]:
        summaries = self.memory.read_recent_summaries()
        if not summaries or len(summaries) < 200:
            return []
        words = re.findall(r'[\u4e00-\u9fff]{2,4}', summaries)
        if not words:
            return []
        ratio = len(set(words)) / len(words)
        if ratio < 0.15:
            return [{"severity": "info", "category": "词汇多样性",
                "description": f"词汇多样性{ratio:.1%}", "suggestion": "增加同义词替换"}]
        return []

    def _check_character_balance(self, ch: int) -> List[Dict]:
        summaries = self.memory.read_recent_summaries()
        if not summaries:
            return []
        last_seen: Dict[str, int] = {}
        for line in summaries.split("\n"):
            m = re.search(r'第(\d+)章', line)
            if m:
                c = int(m.group(1))
                names_m = re.search(r'出场人物[：:]\s*(.+?)(?:\s*\||$)', line)
                if names_m:
                    for name in re.split(r'[,，、]', names_m.group(1)):
                        name = name.strip()
                        if len(name) >= 2:
                            last_seen[name] = max(last_seen.get(name, 0), c)
        issues = []
        for name, last in last_seen.items():
            gap = ch - last
            if gap >= 20:
                issues.append({"severity": "info", "category": "角色出场均衡",
                    "description": f"'{name}'已{gap}章未出场", "suggestion": "确认是否被遗忘"})
        return issues

    def _check_numerical_consistency(self) -> List[Dict]:
        ledger = self.memory.read_resource_ledger()
        if not ledger:
            return []
        issues = []
        for m in re.finditer(r'(\S{2,8}?)\s*[:：]\s*(-\d+)', ledger):
            issues.append({"severity": "warning", "category": "数值体系",
                "description": f"'{m.group(1)}'出现负数：{m.group(2)}", "suggestion": "检查扣减逻辑"})
        return issues


def format_audit_report(issues: List[Dict], ch: int) -> str:
    if not issues:
        return f"  [审计] 第{ch}章：全部通过"
    warnings = [i for i in issues if i["severity"] == "warning"]
    infos = [i for i in issues if i["severity"] == "info"]
    lines = [f"  [审计] 第{ch}章：{len(warnings)}个警告，{len(infos)}个提示"]
    for i in warnings:
        lines.append(f"    ⚠ [{i['category']}] {i['description']}")
        lines.append(f"      → {i['suggestion']}")
    for i in infos[:5]:
        lines.append(f"    ℹ [{i['category']}] {i['description']}")
    if len(infos) > 5:
        lines.append(f"    ℹ ...还有{len(infos) - 5}个提示")
    return "\n".join(lines)


# ============================================================
# ContextBuilder - 精简上下文
# ============================================================

class ContextBuilder:

    STOP_WORDS = {"一个", "这个", "那个", "他们", "她们", "我们", "你们", "什么", "怎么",
                  "为什么", "可以", "没有", "但是", "然后", "因为", "所以", "如果", "虽然",
                  "只是", "已经", "正在", "开始", "发现", "决定", "告诉", "知道", "认为",
                  "觉得", "来到", "看到", "听到", "想到", "说道"}

    @staticmethod
    def _keywords(text: str) -> List[str]:
        cn = re.findall(r'[\u4e00-\u9fff]{2,4}', text)
        en = [w for w in re.findall(r'[A-Za-z]{3,}', text)]
        words = [w for w in cn + en if w.lower() not in ContextBuilder.STOP_WORDS]
        seen = set()
        return [w for w in words if not (w in seen or seen.add(w))][:30]

    @staticmethod
    def filter_relevant(text: str, reference: str, max_items: int = 8) -> str:
        if not text:
            return "（暂无记录）"
        keywords = ContextBuilder._keywords(reference)
        lines = [l.strip() for l in text.strip().split("\n") if l.strip()]
        if not keywords:
            return "\n".join(lines[-max_items:])
        scored = [(sum(1 for kw in keywords if kw in l), l) for l in lines]
        relevant = [(s, l) for s, l in scored if s > 0]
        if relevant:
            relevant.sort(key=lambda x: x[0], reverse=True)
            return "\n".join(l for _, l in relevant[:max_items])
        return "\n".join(lines[-max_items:])

    @staticmethod
    def build_context(chapter_num: int, memory: 'MemorySystem', chapter_outline: str) -> Dict[str, str]:
        return {
            "global_memory": ContextBuilder.filter_relevant(memory.read_global_memory(), chapter_outline, 10),
            "clue_ledger": ContextBuilder.filter_relevant(memory.read_clue_ledger(), chapter_outline, 10),
            "character_matrix": ContextBuilder.filter_relevant(memory.read_character_matrix(), chapter_outline, 8),
            "resource_ledger": ContextBuilder.filter_relevant(memory.read_resource_ledger(), chapter_outline, 5),
            "emotional_arcs": ContextBuilder.filter_relevant(memory.read_emotional_arcs(), chapter_outline, 5),
            "recent_summaries": memory.read_recent_summaries(),
        }


# ============================================================
# Config
# ============================================================

class Config:

    def __init__(self):
        self.project_name = "我的小说"
        self.total_chapters = 500
        self.words_per_chapter = 2500
        self.ai_base_url = ""
        self.ai_api_key = ""
        self.ai_model = ""
        self.writing_style = ""
        self.special_requirements = ""
        self.review_standards = ""
        self.skip_material_check = False
        self.chapter_retry_max = 3
        self.final_review_max_rounds = 3
        self.batch_review_interval = 50
        self.recent_summary_window = 3
        self.transition_word_interval = 1000
        self.max_audit_revise_attempts = 2
        self.base_path = ""
        self.memory_path = ""
        self.chapter_path = ""
        self.review_path = ""
        self.original_path = ""

    def setup_paths(self):
        self.base_path = os.path.expanduser(f"~/Desktop/{self.project_name}")
        self.memory_path = os.path.join(self.base_path, "记忆系统")
        self.chapter_path = os.path.join(self.base_path, "章节")
        self.review_path = os.path.join(self.base_path, "审核记录")
        self.original_path = os.path.join(self.base_path, "原始材料")
        for path in [self.base_path, self.memory_path, self.chapter_path,
                     self.review_path, self.original_path]:
            os.makedirs(path, exist_ok=True)


# ============================================================
# AIInterface
# ============================================================

class AIInterface:

    def __init__(self, config: Config):
        self.config = config
        self.client = OpenAI(api_key=config.ai_api_key, base_url=config.ai_base_url)
        self.model = config.ai_model

    def generate(self, system: str, user: str) -> str:
        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            temperature=0.8,
            max_tokens=12288,
        )
        text = response.choices[0].message.content or ""
        text = re.sub(r'<thinking>.*?</thinking>', '', text, flags=re.DOTALL)
        text = re.sub(r'&lt;think&gt;.*?&lt;/think&gt;', '', text, flags=re.DOTALL)
        text = re.sub(r'$$thinking$$.*?$$/thinking$$', '', text, flags=re.DOTALL)
        return text.strip()


# ============================================================
# MemorySystem（含新台账文件）
# ============================================================

class MemorySystem:

    MEMORY_FILES = [
        "全局记忆.docx", "线索台账.docx", "近期摘要.docx",
        "章节大纲索引.docx", "写作进度.docx",
        "角色信息边界.docx", "资源账本.docx", "情感弧线.docx",
    ]

    def __init__(self, config: Config):
        self.config = config

    def initialize(self, story_bible: str, chapter_outlines: str):
        self._write_file("全局记忆.docx", "（等待AI从故事圣经中提取）")
        self._write_file("线索台账.docx", "（等待AI从故事圣经中提取）")
        self._write_file("近期摘要.docx", "（初始为空，写完第一章后开始填充）")
        self._write_file("章节大纲索引.docx", chapter_outlines)
        self._write_file("写作进度.docx", self._format_progress(0, 1, "写作中"))
        self._write_file("角色信息边界.docx", "（初始为空，写完第一章后开始填充）")
        self._write_file("资源账本.docx", "（初始为空，写完第一章后开始填充）")
        self._write_file("情感弧线.docx", "（初始为空，写完第一章后开始填充）")

    # ---- 读取 ----

    def read_global_memory(self) -> str:
        return self._read_file("全局记忆.docx")

    def read_clue_ledger(self) -> str:
        return self._read_file("线索台账.docx")

    def read_recent_summaries(self) -> str:
        return self._read_file("近期摘要.docx")

    def read_character_matrix(self) -> str:
        return self._read_file("角色信息边界.docx")

    def read_resource_ledger(self) -> str:
        return self._read_file("资源账本.docx")

    def read_emotional_arcs(self) -> str:
        return self._read_file("情感弧线.docx")

    def read_chapter_outline(self, chapter_num: int) -> str:
        content = self._read_file("章节大纲索引.docx")
        pattern = rf"【第{chapter_num}章大纲】\s*\n(.*?)(?=\n【第\d+章大纲】|\Z)"
        match = re.search(pattern, content, re.DOTALL)
        return match.group(1).strip() if match else ""

    def read_progress(self) -> dict:
        content = self._read_file("写作进度.docx")
        progress = {}
        for line in content.strip().split("\n"):
            if "：" in line:
                key, value = line.split("：", 1)
                progress[key.strip("- ").strip()] = value.strip()
        return progress

    # ---- 更新 ----

    def update_progress(self, completed: int, next_chapter: int, status: str, context: str):
        content = self._format_progress(completed, next_chapter, status, context)
        self._write_file("写作进度.docx", content)

    def update_global_memory(self, updates: str):
        if not updates or updates.strip() == "无变化":
            return
        current = self.read_global_memory()
        self._write_file("全局记忆.docx", current + "\n\n" + updates)

    def update_clue_ledger(self, clue_changes: list):
        if not clue_changes:
            return
        current = self.read_clue_ledger()
        changes_text = "\n".join(f"- {c}" for c in clue_changes)
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        if "【最后更新时间】" in current:
            updated = current.replace("【最后更新时间】", f"{changes_text}\n\n【最后更新时间】")
        else:
            updated = current + f"\n\n{changes_text}\n\n【最后更新时间】{now}"
        self._write_file("线索台账.docx", updated)

    def update_recent_summaries(self, chapter_num: int, summary: str):
        current = self.read_recent_summaries()
        lines = [l for l in current.strip().split("\n") if l.strip() and not l.startswith("【")]
        lines.append(f"第{chapter_num}章：{summary}")
        if len(lines) > self.config.recent_summary_window:
            lines = lines[-self.config.recent_summary_window:]
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        content = "\n".join(lines) + f"\n\n【最后更新时间】{now}"
        self._write_file("近期摘要.docx", content)

    def update_character_matrix(self, chapter_num: int, entries: list):
        if not entries:
            return
        current = self.read_character_matrix()
        changes_text = "\n".join(f"- {e}" for e in entries)
        entry = f"\n\n## 第{chapter_num}章更新\n{changes_text}"
        self._write_file("角色信息边界.docx", current + entry)

    def update_resource_ledger(self, chapter_num: int, entries: list):
        if not entries:
            return
        current = self.read_resource_ledger()
        changes_text = "\n".join(f"- {e}" for e in entries)
        entry = f"\n\n## 第{chapter_num}章更新\n{changes_text}"
        self._write_file("资源账本.docx", current + entry)

    def update_emotional_arcs(self, chapter_num: int, entries: list):
        if not entries:
            return
        current = self.read_emotional_arcs()
        changes_text = "\n".join(f"- {e}" for e in entries)
        entry = f"\n\n## 第{chapter_num}章更新\n{changes_text}"
        self._write_file("情感弧线.docx", current + entry)

    # ---- 内部 ----

    def _format_progress(self, completed: int, next_chapter: int, status: str, context: str = "") -> str:
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        return (f"- 已完成：{completed}章\n- 下一章：第{next_chapter}章\n"
                f"- 状态：{status}\n- 最后更新时间：{now}\n- 当前任务上下文：{context}")

    def _write_file(self, filename: str, content: str):
        filepath = os.path.join(self.config.memory_path, filename)
        atomic_write(filepath, content, as_docx=True)

    def _read_file(self, filename: str) -> str:
        filepath = os.path.join(self.config.memory_path, filename)
        if not os.path.exists(filepath):
            return ""
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)


# ============================================================
# DocxHandler
# ============================================================

class DocxHandler:

    def __init__(self, config: Config):
        self.config = config

    def save_chapter(self, chapter_num: int, content: str) -> str:
        filepath = os.path.join(self.config.chapter_path, f"第{chapter_num:03d}章.docx")
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)

        paragraphs = [p.strip() for p in content.strip().split("\n") if p.strip()]
        for i, para_text in enumerate(paragraphs):
            p = doc.add_paragraph()
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = Pt(22)
                run = p.add_run(para_text)
                run.font.name = "宋体"
                run.font.size = Pt(14)
                run.bold = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(24)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(22)
                run = p.add_run(para_text)
                run.font.name = "宋体"
                run.font.size = Pt(12)
        atomic_write(filepath, "", as_docx=False)
        doc.save(filepath)
        return filepath

    def format_check(self, chapter_num: int, content: str) -> Tuple[bool, list]:
        issues = []
        paragraphs = [p.strip() for p in content.split("\n") if p.strip()]
        body_text = "".join(paragraphs[1:]) if len(paragraphs) > 1 else ""
        char_count = len(body_text.replace(" ", ""))
        min_words = self.config.words_per_chapter - 300
        max_words = self.config.words_per_chapter + 300
        if char_count < min_words or char_count > max_words:
            issues.append(f"字数异常：{char_count}字，要求{min_words}-{max_words}字")
        for i, p in enumerate(paragraphs[1:], 2):
            p_len = len(p.replace(" ", ""))
            if p_len < 50:
                issues.append(f"第{i}段过短：{p_len}字")
            elif p_len > 250:
                issues.append(f"第{i}段过长：{p_len}字")
        forbidden = [r"---元数据---", r"---元数据结束---", r"摘要：", r"线索变化："]
        for pattern in forbidden:
            if re.search(pattern, content):
                issues.append(f"正文纯度违规：{pattern}")
        return len(issues) == 0, issues

    def fix_format(self, content: str) -> str:
        content = re.sub(r"\n{3,}", "\n\n", content)
        content = re.sub(r"(?<!\w)\*{1,3}(?!\w)", "", content)
        content = re.sub(r"^#{1,6}\s", "", content, flags=re.MULTILINE)
        content = re.sub(r"^>\s", "", content, flags=re.MULTILINE)
        content = content.replace("```", "")
        return content.strip()

    def save_review_record(self, filename: str, content: str):
        filepath = os.path.join(self.config.review_path, filename)
        doc = Document()
        doc.add_paragraph(content)
        doc.save(filepath)


# ============================================================
# ReviewStandards
# ============================================================

class ReviewStandards:

    HARD_ERRORS = [(r"第\d+章|chapter\s+\d+", "禁止章节号指称")]
    REPORT_TERMS = ["核心动机", "信息落差", "认知共鸣", "锚定效应",
                    "叙事张力", "情感锚点", "角色弧光", "戏剧冲突"]
    TRANSITION_WORDS = ["仿佛", "忽然", "竟然", "猛地", "猛然", "不禁", "宛如"]
    PREACHY_WORDS = ["显然", "毋庸置疑", "不言而喻", "众所周知"]
    SHOCK_PHRASES = ["全场震惊", "众人惊呆", "所有人都愣住了", "一片哗然"]

    @staticmethod
    def check_hard_errors(content: str) -> list:
        issues = []
        for pattern, desc in ReviewStandards.HARD_ERRORS:
            if re.search(pattern, content):
                issues.append(f"硬性错误：{desc}")
        for term in ReviewStandards.REPORT_TERMS:
            if term in content:
                issues.append(f"硬性错误：包含分析报告术语「{term}」")
        md_patterns = [(r"(?<!\w)\*{1,3}(?!\w)", "Markdown星号"),
                       (r"^#{1,6}\s", "Markdown标题"),
                       (r"^>\s", "Markdown引用"), (r"```", "Markdown代码块")]
        for pattern, desc in md_patterns:
            if re.search(pattern, content, re.MULTILINE):
                issues.append(f"硬性错误：包含{desc}")
        return issues

    @staticmethod
    def check_warnings(content: str, transition_interval: int = 1000) -> list:
        issues = []
        char_count = len(content.replace("\n", "").replace(" ", ""))
        transition_count = sum(content.count(w) for w in ReviewStandards.TRANSITION_WORDS)
        max_transitions = max(1, char_count // transition_interval)
        if transition_count > max_transitions:
            issues.append(f"警告：转折词密度过高（{transition_count}次/{char_count}字）")
        for word in ReviewStandards.PREACHY_WORDS:
            if word in content:
                issues.append(f"警告：包含说教词「{word}」")
        for phrase in ReviewStandards.SHOCK_PHRASES:
            if phrase in content:
                issues.append(f"警告：包含集体震惊套话「{phrase}」")
        sentences = re.split(r"[。！？]", content)
        consecutive_le = 0
        for s in sentences:
            if "了" in s:
                consecutive_le += 1
                if consecutive_le >= 6:
                    issues.append("警告：连续6句以上含「了」字")
                    break
            else:
                consecutive_le = 0
        words = re.findall(r"[\u4e00-\u9fff]{2,4}", content)
        word_count = {}
        for w in words:
            word_count[w] = word_count.get(w, 0) + 1
        for w, c in word_count.items():
            if c > 3 and len(w) <= 2:
                issues.append(f"警告：高频词「{w}」出现{c}次")
        return issues


# ============================================================
# PromptBuilder（注入 InkOS 写作方法论）
# ============================================================

class PromptBuilder:

    # InkOS 写作方法论核心（精简版，从 writer-prompts.ts 提取）
    WRITING_METHODOLOGY = """
## 人物塑造铁律
- 人设一致性：角色行为由"过往经历 + 当前利益 + 性格底色"共同驱动
- 人物立体化：核心标签 + 反差细节 = 活人
- 拒绝工具人：配角必须有独立动机和反击能力
- 角色区分度：不同角色的语气、发怒方式、处事模式必须有显著差异

## 信息边界铁律
- 每个角色只能基于已知信息行动和思考
- 角色A没有亲历或被告知事件B，A不能对B做出反应
- 多角色同场景时，每个角色的台词只能包含该角色已知的信息

## 去AI味铁律
- 叙述者永远不得替读者下结论
- 禁止分析报告式语言（核心动机、信息边界、信息落差、核心风险、利益最大化、当前处境）
- 转折/惊讶标记词（仿佛、忽然、竟、竟然、猛地、猛然、不禁、宛如）全篇每3000字最多1次
- 同一体感/意象禁止连续渲染超过两轮
- 禁止"不是……而是……"句式
- 禁止破折号"——"

## 反例→正例对照
情绪描写：
  ✗ 他感到非常愤怒 → ✓ 他捏碎了手中的茶杯，滚烫的茶水流过指缝
  ✗ 她心里很悲伤 → ✓ 她攥紧手机，指节发白
  ✗ 他感到一阵恐惧 → ✓ 他后背的汗毛竖了起来
转折与衔接：
  ✗ 虽然他很强，但是他还是输了 → ✓ 他确实强，可对面那个老东西更脏
  ✗ 然而，事情并没有那么简单 → ✓ 哪有那么便宜的事
  ✗ 因此，他决定采取行动 → ✓ 他站起来，把凳子踢到一边
叙述者姿态：
  ✗ 这一刻，他终于明白了什么是力量 → 删掉，让读者自己感受
  ✗ 显然，对方低估了他 → 只写对方的表情变化
  ✗ 全场为之震惊 → 老陈的烟掉在裤子上，烫得他跳起来

## 六步走人物心理分析
1. 当前处境 2. 核心动机 3. 信息边界 4. 性格过滤 5. 行为选择 6. 情绪外化
禁止跳过步骤直接写行为。

## 人设防崩三问
1. "为什么这么做？"——必须有驱动
2. "符合之前的人设吗？"
3. "读者会觉得突兀吗？"

## 创作宪法（14条）
1. Show don't tell 2. 价值观像盐溶于汤 3. 行动立于三条腿：过往经历、当前利益、性格底色
4. 配角都有自己的账本 5. 节奏即呼吸 6. 每章结尾有钩子 7. 全员智商在线
8. 后世梗符合年代 9. 时间线不能错 10. 日常七成必须成为伏笔 11. 关系改变要事件驱动
12. 人设前后一致 13. 重要剧情用场景不用总结 14. 拒绝流水账

## 代入感六支柱
1. 基础信息标签化 2. 可视化熟悉感 3. 共鸣分两层（认知+情绪）
4. 欲望两条腿（基础+主动） 5. 五感钩子 6. 人设活化（标签+反差）

## 黄金三章纪律
第一章：开篇直接进入冲突，第一段必须有动作或对话，最多3个角色
第二章：核心优势通过具体事件初现，第一个小爽点
第三章：具体可衡量的短期目标，章尾强钩子
"""

    def __init__(self, config: Config):
        self.config = config

    def build_material_confirmation_prompt(
        self, story_bible: str, chapter_outlines: str,
        total_chapters: int, words_per_chapter: int,
    ) -> Tuple[str, str]:
        system = (
            '你是一个小说项目审核AI。启动前检查材料是否完整、有无矛盾、有无模糊不清。\n\n'
            '【检查项】\n1. 故事圣经是否包含：世界观设定、核心角色档案、主线/支线剧情框架\n'
            '2. 章节大纲是否覆盖全部章节\n3. 故事圣经和大纲之间是否有逻辑矛盾\n'
            '4. 是否有模糊不清需要澄清的地方\n5. 章节数量是否匹配\n\n'
            '【输出格式】\n{"passed": true/false, "issues": ["问题描述"], "suggestions": ["建议"]}'
        )
        user = (f"【总章节数】{total_chapters}\n【每章目标字数】{words_per_chapter}\n\n"
                f"【故事圣经】\n{story_bible}\n\n【章节大纲】\n{chapter_outlines}")
        return system, user

    def build_writing_prompt(
        self, chapter_num: int, chapter_outline: str, context: Dict[str, str],
        feedback: str = None,
    ) -> Tuple[str, str]:
        style = f"\n【写作风格要求】\n{self.config.writing_style}" if self.config.writing_style else ""
        special = f"\n【特殊要求】\n{self.config.special_requirements}" if self.config.special_requirements else ""
        custom_review = f"\n【自定义审核标准】\n{self.config.review_standards}" if self.config.review_standards else ""

        system = (
            f'你是一个小说写作AI。根据提供的材料写一章正文。\n\n'
            f'【写作要求】\n'
            f'1. 创作约{self.config.words_per_chapter}字正文\n'
            f'2. 不得与记忆系统中的任何信息矛盾\n'
            f'3. 自然收尾，留白或场景转换，禁止总结性段落和AI式结尾\n'
            f'4. 第一行必须是章节标题，格式为「第X章 标题」\n'
            f'{style}{special}{custom_review}\n\n'
            f'{self.WRITING_METHODOLOGY}\n\n'
            f'【违禁规则 - 硬性】\n'
            f'- 禁止章节号指称（第X章、chapter X）\n'
            f'- 禁止分析报告术语（核心动机、信息落差、认知共鸣、锚定效应等）\n'
            f'- 禁止未设定的人物/地名\n'
            f'- 禁止真实地名（除非大纲明确）\n'
            f'- 禁止违背现实逻辑的情节\n'
            f'- 禁止Markdown格式符号（*、#、-、_、```、> 等）\n\n'
            f'【违禁规则 - 警告】\n'
            f'- 转折词密度：每{self.config.transition_word_interval}字不超过1次\n'
            f'- 高疲劳词：同词每章只出现1次\n'
            f'- 禁止元叙事/编剧旁白\n'
            f'- 禁止说教词（显然、毋庸置疑、不言而喻、众所周知）\n'
            f'- 禁止集体震惊套话（全场震惊、众人惊呆）\n'
            f'- 禁止连续6句以上含「了」字\n'
            f'- 段落控制在50-250字\n\n'
            f'【信息边界】\n'
            f'角色只能基于自己已知的信息行动和思考。违反信息边界是严重错误。\n\n'
            f'【输出格式 - 严格遵守】\n'
            f'正文结束后，另起一段，严格按以下格式输出元数据：\n\n'
            f'---元数据---\n'
            f'摘要：（100字以内的本章内容摘要）\n'
            f'线索变化：\n'
            f'- （线索名称）：从（旧状态）变为（新状态）\n'
            f'角色信息边界：\n'
            f'- 角色名 | 本章获知的信息 | 本章仍不知道的信息\n'
            f'资源变动：\n'
            f'- 物品名 | 归属/状态变化\n'
            f'情感变化：\n'
            f'- 角色名 | 情绪变化\n'
            f'建议更新：（如有需要更新全局记忆的内容，没有则写「无」）\n'
            f'---元数据结束---\n\n'
            f'注意：元数据仅用于系统更新记忆，不会写入正文。'
        )

        user = f"【世界观规则】\n{context.get('global_memory', '')}\n\n"
        user += f"【线索台账】\n{context.get('clue_ledger', '')}\n\n"
        user += f"【角色信息边界】\n{context.get('character_matrix', '')}\n\n"
        user += f"【资源账本】\n{context.get('resource_ledger', '')}\n\n"
        user += f"【情感弧线】\n{context.get('emotional_arcs', '')}\n\n"
        user += f"【近期摘要】\n{context.get('recent_summaries', '')}\n\n"
        user += f"【第{chapter_num}章大纲】\n{chapter_outline}\n\n"
        user += f"请创作第{chapter_num}章正文。"
        if feedback:
            user += f"\n\n【上一次审核反馈，请务必修正】\n{feedback}"
        return system, user

    def build_review_prompt(
        self, content: str, clue_ledger: str,
        recent_summaries: str, chapter_num: int,
    ) -> Tuple[str, str]:
        custom_review = f"\n【自定义审核标准】\n{self.config.review_standards}" if self.config.review_standards else ""
        system = (
            '你是一个小说审核AI。严格检查以下章节内容。\n\n'
            '【检查项】\n1. 与近期摘要中上一章结尾是否无缝衔接\n'
            '2. 本章内容是否与线索台账矛盾\n3. 人物状态是否与全局记忆一致\n'
            '4. 情节是否符合大纲要求\n5. 是否存在硬性违禁项\n6. 是否存在警告项\n'
            f'{custom_review}\n\n'
            '【severity严重程度说明】\n- hard：硬性错误，必须重写\n- warning：警告项，需要修改\n- none：无问题\n\n'
            '【输出格式】\n{"passed": true/false, "issues": ["问题描述"], "severity": "hard/warning/none"}'
        )
        user = (f"【线索台账】\n{clue_ledger}\n\n【近期摘要】\n{recent_summaries}\n\n"
                f"【待审核章节 - 第{chapter_num}章】\n{content}")
        return system, user

    def build_memory_init_prompt(self, story_bible: str) -> Tuple[str, str]:
        system = (
            '你是一个小说设定分析AI。从故事圣经中提取信息，严格按JSON格式输出：\n\n'
            '{\n'
            '  "global_memory": "世界观核心规则（不超过500字）\\n\\n核心角色档案（姓名、年龄、身份、性格、关键特征、初始状态）",\n'
            '  "clue_ledger": "角色状态追踪\\n物品道具追踪\\n伏笔追踪\\n时间线\\n\\n【最后更新时间】初始化"\n'
            '}'
        )
        user = f"【故事圣经】\n{story_bible}"
        return system, user

    def build_chapter_analysis_prompt(self, chapter_content: str, chapter_num: int) -> Tuple[str, str]:
        system = (
            '你是一个小说分析AI。分析章节内容，提取关键信息，严格按JSON格式输出：\n'
            '{"summary": "100字以内摘要", '
            '"character_states": ["角色当前状态"], '
            '"clue_changes": ["线索变化"], '
            '"character_matrix": ["角色名 | 获知信息 | 未知信息"], '
            '"resource_changes": ["物品名 | 变化"], '
            '"emotional_changes": ["角色名 | 情绪变化"]}'
        )
        user = f"【第{chapter_num}章内容】\n{chapter_content}"
        return system, user


# ============================================================
# NovelForge - 主进程
# ============================================================

class NovelForge:

    def __init__(self, config: Config):
        self.config = config
        self.ai = AIInterface(config)
        self.memory = MemorySystem(config)
        self.docx = DocxHandler(config)
        self.prompts = PromptBuilder(config)
        self.audit = AuditSystem(self.memory, config)

    def start(self, mode: str, **kwargs):
        if mode == "new":
            self._start_new(**kwargs)
        elif mode == "resume":
            self._start_resume()
        elif mode == "import":
            self._start_import(**kwargs)

    # ---- 启动模式 ----

    def _confirm_materials(self, story_bible: str, chapter_outlines: str) -> Tuple[bool, list]:
        print("[NovelForge] 正在确认材料完整性...")
        system, user = self.prompts.build_material_confirmation_prompt(
            story_bible, chapter_outlines, self.config.total_chapters, self.config.words_per_chapter)
        try:
            response = self.ai.generate(system, user)
            result = self._parse_json(response)
            if result:
                passed = result.get("passed", False)
                issues = result.get("issues", [])
                if passed:
                    print("[NovelForge] 材料确认通过")
                    return True, []
                else:
                    print(f"[NovelForge] 材料确认未通过，{len(issues)}个问题")
                    for i, issue in enumerate(issues, 1):
                        print(f"  {i}. {issue}")
                    return False, issues
            else:
                print("[NovelForge] 材料确认AI返回格式异常，跳过确认")
                return True, []
        except Exception as e:
            print(f"[NovelForge] 材料确认失败：{e}，跳过确认")
            return True, []

    def _start_new(self, story_bible: str, chapter_outlines: str):
        print(f"[NovelForge] v2.3.0 项目：{self.config.project_name}")
        print(f"[NovelForge] 总章数：{self.config.total_chapters}，每章{self.config.words_per_chapter}字")

        if not self.config.skip_material_check:
            passed, issues = self._confirm_materials(story_bible, chapter_outlines)
            if not passed:
                print("\n[NovelForge] 材料有问题，请补充后重新启动")
                for i, issue in enumerate(issues, 1):
                    print(f"  {i}. {issue}")
                print("补充完毕后重新调用 start('new')，或设置 skip_material_check=True 跳过")
                return
        else:
            print("[NovelForge] 已跳过材料确认")

        self.config.setup_paths()
        self._save_original("故事圣经.docx", story_bible)
        self._save_original("章节大纲.docx", chapter_outlines)

        print("[NovelForge] AI正在分析故事圣经，生成记忆系统...")
        system, user = self.prompts.build_memory_init_prompt(story_bible)
        ai_response = self.ai.generate(system, user)
        memory_data = self._parse_json(ai_response)

        self.memory.initialize(story_bible, chapter_outlines)
        if memory_data:
            if "global_memory" in memory_data:
                self.memory._write_file("全局记忆.docx", memory_data["global_memory"])
            if "clue_ledger" in memory_data:
                self.memory._write_file("线索台账.docx", memory_data["clue_ledger"])

        print("[NovelForge] 记忆系统初始化完成")
        print("[NovelForge] 开始自动写作循环")
        self._main_loop(start_chapter=1)

    def _start_resume(self):
        print("[NovelForge] 断点续写")
        self.config.setup_paths()
        progress = self.memory.read_progress()
        next_chapter = int(progress.get("下一章", "第1章").replace("第", "").replace("章", ""))
        print(f"[NovelForge] 从第{next_chapter}章继续")
        self._main_loop(start_chapter=next_chapter)

    def _start_import(self, story_bible: str, chapter_outlines: str,
                      existing_chapters: dict, start_chapter: int):
        print(f"[NovelForge] 导入模式：从第{start_chapter}章开始续写")
        if not self.config.skip_material_check:
            passed, issues = self._confirm_materials(story_bible, chapter_outlines)
            if not passed:
                print("\n[NovelForge] 材料有问题，请补充后重新启动")
                return
        self.config.setup_paths()
        self._save_original("故事圣经.docx", story_bible)
        self._save_original("章节大纲.docx", chapter_outlines)
        self.memory.initialize(story_bible, chapter_outlines)

        total_import = start_chapter - 1
        print(f"[NovelForge] 开始逐章导入，共{total_import}章")
        for i in range(1, start_chapter):
            chapter_content = existing_chapters.get(i, "")
            if not chapter_content:
                print(f"[NovelForge] 警告：第{i}章内容缺失，跳过")
                continue
            print(f"[NovelForge] 导入第{i}章...")
            system, user = self.prompts.build_chapter_analysis_prompt(chapter_content, i)
            analysis = self.ai.generate(system, user)
            analysis_data = self._parse_json(analysis)
            if analysis_data:
                if "summary" in analysis_data:
                    self.memory.update_recent_summaries(i, analysis_data["summary"])
                if "clue_changes" in analysis_data:
                    self.memory.update_clue_ledger(analysis_data["clue_changes"])
                if "character_matrix" in analysis_data:
                    self.memory.update_character_matrix(i, analysis_data["character_matrix"])
                if "resource_changes" in analysis_data:
                    self.memory.update_resource_ledger(i, analysis_data["resource_changes"])
                if "emotional_changes" in analysis_data:
                    self.memory.update_emotional_arcs(i, analysis_data["emotional_changes"])
            self.docx.save_chapter(i, chapter_content)

        self.memory.update_progress(total_import, start_chapter, "写作中",
                                    f"导入完成，即将开始第{start_chapter}章写作")
        print(f"[NovelForge] 导入完成，开始写作第{start_chapter}章")
        self._main_loop(start_chapter=start_chapter)

    # ---- 主循环 ----

    def _main_loop(self, start_chapter: int):
        for chapter_num in range(start_chapter, self.config.total_chapters + 1):
            print(f"\n{'=' * 60}")
            print(f"[NovelForge] 第{chapter_num}章")
            print(f"{'=' * 60}")

            success = self._write_and_save_chapter(chapter_num)
            if not success:
                print(f"[NovelForge] 第{chapter_num}章处理失败，流程暂停")
                self.memory.update_progress(chapter_num - 1, chapter_num, "暂停-需要人工介入",
                                            f"第{chapter_num}章异常处理全部失败")
                return

            if chapter_num % self.config.batch_review_interval == 0:
                print(f"[NovelForge] 触发阶段审核")
                self._batch_review(chapter_num)

        print(f"\n[NovelForge] 全部{self.config.total_chapters}章完成，开始终审")
        self._final_review()

    # ---- 单章流程 ----

    def _write_and_save_chapter(self, chapter_num: int) -> bool:
        chapter_outline = self.memory.read_chapter_outline(chapter_num)

        # 精简上下文（避免 token 爆炸）
        context = ContextBuilder.build_context(chapter_num, self.memory, chapter_outline)

        # 写作
        print(f"[写作] 调用AI生成第{chapter_num}章...")
        content, metadata = self._call_ai_write(chapter_num, chapter_outline, context)
        if not content:
            return False

        if metadata is None:
            print(f"[截断] 第{chapter_num}章输出被截断，触发重写")

        # spot-fix
        content, fixes = PostWriteValidator.spot_fix(content)
        if fixes:
            print(f"[Spot-fix] {', '.join(fixes)}")

        # AI 痕迹检测
        ai_issues = AITellDetector.analyze(content)
        ai_warnings = [i for i in ai_issues if i["severity"] == "warning"]
        if ai_warnings:
            print(f"[AI检测] {len(ai_warnings)}个警告")
            for i in ai_warnings:
                print(f"  ⚠ [{i['category']}] {i['description']}")

        # 审核
        clue_ledger = self.memory.read_clue_ledger()
        recent_summaries = self.memory.read_recent_summaries()

        print(f"[审核] 审核第{chapter_num}章...")
        passed, issues, severity = self._call_ai_review(content, clue_ledger, recent_summaries, chapter_num)

        if metadata is None:
            passed = False
            if "截断：缺少元数据" not in str(issues):
                issues.append("截断：缺少元数据，输出可能被token限制截断")
            severity = "hard"

        # 重试循环
        retry_count = 0
        while not passed and retry_count < self.config.chapter_retry_max:
            retry_count += 1
            print(f"[审核] 未通过（severity={severity}，第{retry_count}次），重写...")
            feedback = "\n".join(issues)
            content, metadata = self._call_ai_write(chapter_num, chapter_outline, context, feedback=feedback)
            if content:
                content, _ = PostWriteValidator.spot_fix(content)
                if metadata is None:
                    passed = False
                    issues = ["截断：缺少元数据"]
                    severity = "hard"
                else:
                    passed, issues, severity = self._call_ai_review(content, clue_ledger, recent_summaries, chapter_num)

        # 异常处理
        if not passed:
            print(f"[异常] 第{chapter_num}章审核3次未通过，进入异常处理")
            content, metadata = self._handle_failure(chapter_num, chapter_outline, context)
            if not content:
                return False

        # 18 维度审计
        audit_issues = self.audit.run_audit(chapter_num, content)
        if audit_issues:
            print(format_audit_report(audit_issues, chapter_num))
            fixable = self.audit.get_fixable_warnings(audit_issues)
            if fixable:
                content = self._audit_revise_loop(chapter_num, content, chapter_outline, context, fixable)
        else:
            print(f"  [审计] 通过")

        # 保存
        print(f"[保存] 保存第{chapter_num}章.docx")
        self.docx.save_chapter(chapter_num, content)

        format_ok, format_issues = self.docx.format_check(chapter_num, content)
        if not format_ok:
            print(f"[格式] 格式检查未通过，自动修复：{format_issues}")
            content = self.docx.fix_format(content)
            self.docx.save_chapter(chapter_num, content)

        # 更新记忆系统
        if metadata:
            self.memory.update_recent_summaries(chapter_num, metadata.get("summary", ""))
            self.memory.update_clue_ledger(metadata.get("clue_changes", []))
            self.memory.update_character_matrix(chapter_num, metadata.get("character_matrix", []))
            self.memory.update_resource_ledger(chapter_num, metadata.get("resource_changes", []))
            self.memory.update_emotional_arcs(chapter_num, metadata.get("emotional_changes", []))
            if metadata.get("suggested_updates"):
                self.memory.update_global_memory("\n".join(metadata["suggested_updates"]))

        self.memory.update_progress(chapter_num, chapter_num + 1, "写作中",
                                    f"第{chapter_num}章已完成")
        print(f"[完成] 第{chapter_num}章完成")
        return True

    # ---- 审计→修订循环 ----

    def _audit_revise_loop(self, chapter_num: int, content: str,
                           chapter_outline: str, context: Dict[str, str],
                           fixable: list) -> str:
        for attempt in range(self.config.max_audit_revise_attempts):
            print(f"  [审计修订] 第{attempt + 1}次（{len(fixable)}个问题）...")
            feedback = "\n".join(f"[{i['category']}] {i['suggestion']}" for i in fixable)
            revised_content, revised_meta = self._call_ai_write(
                chapter_num, chapter_outline, context, feedback=feedback)
            if revised_content:
                revised_content, _ = PostWriteValidator.spot_fix(revised_content)
                re_issues = self.audit.run_audit(chapter_num, revised_content)
                re_fixable = self.audit.get_fixable_warnings(re_issues)
                if len(re_fixable) < len(fixable):
                    content = revised_content
                    fixable = re_fixable
                    if not re_fixable:
                        print(f"  [审计修订] 已修复")
                        break
                else:
                    print(f"  [审计修订] 无改善，保留当前版本")
                    break
        return content

    # ---- AI 调用 ----

    def _call_ai_write(self, chapter_num: int, chapter_outline: str,
                       context: Dict[str, str], feedback: str = None) -> Tuple[Optional[str], Optional[dict]]:
        try:
            system, user = self.prompts.build_writing_prompt(
                chapter_num, chapter_outline, context, feedback=feedback)
            response = self.ai.generate(system, user)
            content, metadata = self._split_content_and_metadata(response)
            return content, metadata
        except Exception as e:
            print(f"[错误] AI写作调用失败：{e}")
            return None, None

    def _call_ai_review(self, content: str, clue_ledger: str,
                        recent_summaries: str, chapter_num: int) -> Tuple[bool, list, str]:
        hard_issues = ReviewStandards.check_hard_errors(content)
        warning_issues = ReviewStandards.check_warnings(content, self.config.transition_word_interval)
        all_local = hard_issues + warning_issues
        if hard_issues:
            return False, all_local, "hard"
        try:
            system, user = self.prompts.build_review_prompt(content, clue_ledger, recent_summaries, chapter_num)
            response = self.ai.generate(system, user)
            review_data = self._parse_json(response)
            if review_data:
                passed = review_data.get("passed", False)
                ai_issues = review_data.get("issues", [])
                severity = review_data.get("severity", "none")
                if severity == "hard":
                    passed = False
                return passed, all_local + ai_issues, severity
            else:
                local_severity = "warning" if warning_issues else "none"
                return len(hard_issues) == 0, all_local, local_severity
        except Exception as e:
            print(f"[错误] AI审核调用失败：{e}")
            local_severity = "warning" if warning_issues else "none"
            return len(hard_issues) == 0, all_local, local_severity

    # ---- 异常处理（4 轮）----

    def _handle_failure(self, chapter_num: int, chapter_outline: str,
                        context: Dict[str, str]) -> Tuple[Optional[str], Optional[dict]]:
        # 第一轮：原始重试
        print("[异常处理] 第一轮：原始重试")
        for i in range(3):
            print(f"[异常处理] 重试 {i + 1}/3")
            content, metadata = self._call_ai_write(chapter_num, chapter_outline, context)
            if content and metadata is not None:
                passed, _, _ = self._call_ai_review(
                    content, self.memory.read_clue_ledger(),
                    self.memory.read_recent_summaries(), chapter_num)
                if passed:
                    return content, metadata

        # 第二轮：精简内容重试
        print("[异常处理] 第二轮：精简内容重试")
        minimal_context = {k: v[:300] if isinstance(v, str) else v for k, v in context.items()}
        minimal_outline = chapter_outline[:500]
        content, metadata = self._call_ai_write(chapter_num, minimal_outline, minimal_context)
        if content and metadata is not None:
            return content, metadata

        # 第三轮：等待冷却后重试
        print("[异常处理] 第三轮：等待5分钟后重试")
        time.sleep(300)
        content, metadata = self._call_ai_write(chapter_num, chapter_outline, context)
        if content and metadata is not None:
            passed, _, _ = self._call_ai_review(
                content, self.memory.read_clue_ledger(),
                self.memory.read_recent_summaries(), chapter_num)
            if passed:
                return content, metadata

        # 第四轮：暂停
        print("[异常处理] 第四轮：全部失败，暂停")
        failure_record = (
            f"第{chapter_num}章异常排查记录\n"
            f"时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"第一轮：原始重试3次，全部失败\n"
            f"第二轮：精简内容重试，失败\n"
            f"第三轮：等待5分钟冷却后重试，失败\n"
            f"第四轮：暂停，等待人工介入"
        )
        self.docx.save_review_record(f"第{chapter_num:03d}章_异常排查.docx", failure_record)
        return None, None

    # ---- 阶段审核 ----

    def _batch_review(self, up_to_chapter: int):
        print(f"\n[阶段审核] 对第1-{up_to_chapter}章进行台账审查")
        clue_ledger = self.memory.read_clue_ledger()
        global_memory = self.memory.read_global_memory()
        recent_summaries = self.memory.read_recent_summaries()

        system = (
            '你是一个小说审核编辑。严格审查线索台账，检查：\n'
            '1. 每条线索的演变是否自洽\n2. 人物状态是否合理\n'
            '3. 伏笔是否有遗忘\n4. 时间线是否混乱\n\n'
            '只有发现可疑项时，才需要回溯相关章节原文核实。\n\n'
            '{"passed": true/false, "issues": ["问题描述"], '
            '"fix_suggestions": ["修复建议"], '
            '"chapters_to_rewrite": [需要重写的章节号列表]}'
        )
        user = (f"【全局记忆】\n{global_memory}\n\n【线索台账】\n{clue_ledger}\n\n"
                f"【近期摘要】\n{recent_summaries}")

        try:
            response = self.ai.generate(system, user)
            review_data = self._parse_json(response)
            if review_data and not review_data.get("passed", True):
                issues = review_data.get("issues", [])
                fixes = review_data.get("fix_suggestions", [])
                chapters_to_rewrite = review_data.get("chapters_to_rewrite", [])
                report = (
                    f"第{up_to_chapter}章阶段审核报告\n"
                    f"时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                    f"问题：\n" + "\n".join(f"- {i}" for i in issues) +
                    f"\n修复建议：\n" + "\n".join(f"- {f}" for f in fixes) +
                    f"\n需要重写的章节：{chapters_to_rewrite}"
                )
                self.docx.save_review_record(f"第{up_to_chapter}章阶段审核.docx", report)
                print(f"[阶段审核] 发现{len(issues)}个问题，需要重写{len(chapters_to_rewrite)}章")
                for chapter_num in chapters_to_rewrite:
                    print(f"[阶段审核] 重写第{chapter_num}章...")
                    self._rewrite_chapter(chapter_num)
                print("[阶段审核] 所有问题已修复，继续写作")
            else:
                print("[阶段审核] 通过")
        except Exception as e:
            print(f"[阶段审核] 审核调用失败：{e}")

    def _rewrite_chapter(self, chapter_num: int) -> bool:
        chapter_outline = self.memory.read_chapter_outline(chapter_num)
        context = ContextBuilder.build_context(chapter_num, self.memory, chapter_outline)
        content, metadata = self._call_ai_write(chapter_num, chapter_outline, context)
        if not content or metadata is None:
            print(f"[重写] 第{chapter_num}章重写失败")
            return False
        content, _ = PostWriteValidator.spot_fix(content)
        self.docx.save_chapter(chapter_num, content)
        format_ok, format_issues = self.docx.format_check(chapter_num, content)
        if not format_ok:
            content = self.docx.fix_format(content)
            self.docx.save_chapter(chapter_num, content)
        if metadata:
            self.memory.update_recent_summaries(chapter_num, metadata.get("summary", ""))
            self.memory.update_clue_ledger(metadata.get("clue_changes", []))
            self.memory.update_character_matrix(chapter_num, metadata.get("character_matrix", []))
            self.memory.update_resource_ledger(chapter_num, metadata.get("resource_changes", []))
            self.memory.update_emotional_arcs(chapter_num, metadata.get("emotional_changes", []))
            if metadata.get("suggested_updates"):
                self.memory.update_global_memory("\n".join(metadata["suggested_updates"]))
        print(f"[重写] 第{chapter_num}章重写完成")
        return True

    # ---- 终审 ----

    def _final_review(self):
        print("[终审] 开始最终审查")
        self._final_review_loop(remaining_rounds=self.config.final_review_max_rounds)

    def _final_review_loop(self, remaining_rounds: int):
        if remaining_rounds <= 0:
            print("[终审] 已重审3轮仍有问题，暂停等待人工介入")
            self.memory.update_progress(
                self.config.total_chapters, self.config.total_chapters,
                "暂停-终审未通过", "终审3轮仍有问题")
            return

        clue_ledger = self.memory.read_clue_ledger()
        global_memory = self.memory.read_global_memory()
        recent_summaries = self.memory.read_recent_summaries()

        system = (
            '你是一个小说终审编辑。对整部小说的线索台账做最终通查。\n'
            '检查所有线索是否完整收束、人物弧线是否完成、是否有遗留问题。\n\n'
            '只有发现可疑项时，才需要回溯相关章节原文核实。\n\n'
            '{"passed": true/false, "issues": ["问题描述"], '
            '"overall_quality": "评价", '
            '"chapters_to_rewrite": [需要重写的章节号列表]}'
        )
        user = (f"【全局记忆】\n{global_memory}\n\n【线索台账】\n{clue_ledger}\n\n"
                f"【近期摘要】\n{recent_summaries}")

        try:
            response = self.ai.generate(system, user)
            review_data = self._parse_json(response)
            if review_data:
                passed = review_data.get("passed", False)
                issues = review_data.get("issues", [])
                chapters_to_rewrite = review_data.get("chapters_to_rewrite", [])

                if not passed:
                    round_num = self.config.final_review_max_rounds - remaining_rounds + 1
                    report = (
                        f"终审报告（第{round_num}轮）\n"
                        f"时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                        f"总章数：{self.config.total_chapters}\n"
                        f"通过：False\n"
                        f"问题：\n" + "\n".join(f"- {i}" for i in issues) +
                        f"\n整体评价：{review_data.get('overall_quality', '')}"
                    )
                    self.docx.save_review_record(f"终审报告_第{round_num}轮.docx", report)
                    print(f"[终审] 第{round_num}轮：{len(issues)}个问题，需要重写{len(chapters_to_rewrite)}章")
                    for chapter_num in chapters_to_rewrite:
                        print(f"[终审] 重写第{chapter_num}章...")
                        self._rewrite_chapter(chapter_num)
                    print("[终审] 问题已修复，重新终审...")
                    self._final_review_loop(remaining_rounds - 1)
                    return
                else:
                    report = (
                        f"终审报告\n"
                        f"时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                        f"总章数：{self.config.total_chapters}\n"
                        f"通过：True\n"
                        f"问题：无\n"
                        f"整体评价：{review_data.get('overall_quality', '')}"
                    )
                    self.docx.save_review_record("终审报告.docx", report)
                    print("[终审] 通过")
        except Exception as e:
            print(f"[终审] 终审调用失败：{e}")
            return

        self.memory.update_progress(
            self.config.total_chapters, self.config.total_chapters,
            "已完成", f"全部{self.config.total_chapters}章写作完成，终审完成")
        print(f"\n[NovelForge] {self.config.project_name} 全部完成！")

    # ---- 工具方法 ----

    def _split_content_and_metadata(self, response: str) -> Tuple[str, Optional[dict]]:
        meta_start = response.find("---元数据---")
        meta_end = response.find("---元数据结束---")

        if meta_start != -1 and meta_end != -1:
            content = response[:meta_start].strip()
            meta_text = response[meta_start + len("---元数据---"):meta_end].strip()

            metadata = {
                "summary": "", "clue_changes": [], "suggested_updates": [],
                "character_matrix": [], "resource_changes": [], "emotional_changes": [],
            }
            current_section = None
            for line in meta_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("摘要："):
                    metadata["summary"] = line.replace("摘要：", "").strip()
                    current_section = "summary"
                elif line.startswith("线索变化："):
                    current_section = "clues"
                elif line.startswith("角色信息边界："):
                    current_section = "matrix"
                elif line.startswith("资源变动："):
                    current_section = "resources"
                elif line.startswith("情感变化："):
                    current_section = "emotions"
                elif line.startswith("建议更新："):
                    value = line.replace("建议更新：", "").strip()
                    if value and value != "无":
                        metadata["suggested_updates"].append(value)
                    current_section = "updates"
                elif line.startswith("- "):
                    item = line[2:].strip()
                    if current_section == "clues":
                        metadata["clue_changes"].append(item)
                    elif current_section == "matrix":
                        metadata["character_matrix"].append(item)
                    elif current_section == "resources":
                        metadata["resource_changes"].append(item)
                    elif current_section == "emotions":
                        metadata["emotional_changes"].append(item)

            return content, metadata

        return response.strip(), None

    def _parse_json(self, text: str) -> Optional[dict]:
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass
        match = re.search(r"```(?:json)?\s*\n?(.*?)\n?```", text, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(1))
            except json.JSONDecodeError:
                pass
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                return json.loads(text[start:end + 1])
            except json.JSONDecodeError:
                pass
        return None

    def _save_original(self, filename: str, content: str):
        filepath = os.path.join(self.config.original_path, filename)
        doc = Document()
        doc.add_paragraph(content)
        doc.save(filepath)


# ============================================================
# 入口
# ============================================================

def main():
    print("=" * 60)
    print("  NovelForge v2.3.0 - 长篇小说全自动写作系统")
    print("=" * 60)

    config = Config()

    print("\n请配置AI接口：")
    config.ai_base_url = input("API Base URL: ").strip()
    config.ai_api_key = input("API Key: ").strip()
    config.ai_model = input("Model Name: ").strip()

    print("\n请配置项目：")
    config.project_name = input("项目名称（小说名）: ").strip()
    config.total_chapters = int(input("总章节数: ").strip())
    words_input = input("每章目标字数（默认2500）: ").strip()
    config.words_per_chapter = int(words_input) if words_input else 2500

    print("\n可选材料（直接回车跳过）：")
    config.writing_style = input("写作风格要求: ").strip()
    config.special_requirements = input("其他特殊要求: ").strip()
    config.review_standards = input("自定义审核标准: ").strip()

    skip_input = input("跳过材料确认？（y/n，默认n）: ").strip()
    config.skip_material_check = skip_input.lower() == "y"

    forge = NovelForge(config)

    print("\n请选择模式：")
    print("1. 开始写（全新写作）")
    print("2. 继续写（断点续写）")
    print("3. 从第X章开始续写（导入模式）")

    choice = input("\n输入选项编号: ").strip()

    if choice == "1":
        print("\n请提供故事圣经（输入文件路径或粘贴内容，输入END结束）：")
        story_bible = _read_multiline()
        print("\n请提供章节大纲（输入文件路径或粘贴内容，输入END结束）：")
        chapter_outlines = _read_multiline()
        forge.start("new", story_bible=story_bible, chapter_outlines=chapter_outlines)

    elif choice == "2":
        forge.start("resume")

    elif choice == "3":
        start_chapter = int(input("从第几章开始续写: ").strip())
        print("\n请提供故事圣经（输入文件路径或粘贴内容，输入END结束）：")
        story_bible = _read_multiline()
        print("\n请提供章节大纲（输入文件路径或粘贴内容，输入END结束）：")
        chapter_outlines = _read_multiline()

        existing_chapters = {}
        print("\n请提供已完成的章节内容（格式：先输入章节数字，再输入内容，输入0结束）")
        while True:
            num = input("\n章节号（输入0结束）: ").strip()
            if num == "0":
                break
            print(f"请输入第{num}章内容（输入END结束）：")
            content = _read_multiline()
            existing_chapters[int(num)] = content

        forge.start("import", story_bible=story_bible,
                    chapter_outlines=chapter_outlines,
                    existing_chapters=existing_chapters,
                    start_chapter=start_chapter)


def _read_multiline() -> str:
    first_line = input().strip()
    if first_line.upper() == 'END':
        return ""
    if first_line.lower().endswith('.docx'):
        parts = first_line.split()
        all_text = []
        for part in parts:
            if part.lower().endswith('.docx'):
                doc = Document(part.strip())
                all_text.append("\n".join(p.text for p in doc.paragraphs))
        return "\n\n".join(all_text)
    lines = [first_line]
    while True:
        line = input()
        if line.strip() == "END":
            break
        lines.append(line)
    return "\n".join(lines)


if __name__ == "__main__":
    main()
